#!/usr/bin/env python3
"""
Briefing tool — MVP (Loops 1 & 2)
=================================

Take a client brief in any format (Word / PDF / text / scraps) and produce:
  * Loop 1 — Parse/Ingest: a faithful structured capture + win-rules, with a
    no-loss ledger proving nothing was dropped. NO research, NO RAG.
  * Loop 2 — First-round brief: shape the capture into an agency brief
    (problem, objective, audience, scope) + the open questions to ask first.

Both loops are written to plain output files (JSON + a markdown one-pager). No
database, no persistent memory in the MVP.

Pipeline:  ingest -> segment -> extract -> [Loop 1] -> review
           -> [Loop 2 shaping] -> review -> render

Usage:
    python parse_brief.py samples/messy_brief_sample.txt \
        --client "Northwind Motors" --project "Moving People"
    python parse_brief.py /path/to/vw_brief.pdf
"""

from __future__ import annotations

import argparse
import base64
import datetime as dt
import json
import os
import re
import shutil
import subprocess
import sys
import urllib.request
import urllib.error
from pathlib import Path

def _load_env_file(path: Path) -> None:
    """Load KEY=VALUE lines from .env into os.environ (without overriding existing).
    Dependency-free fallback — python-dotenv isn't installed in the `claws` env, and
    without this a mis-sourced shell (`. .env` vs `. ./.env`) silently drops all API
    keys and the whole pipeline falls back to heuristic mode. Never let that be silent."""
    try:
        for line in path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            k, v = k.strip(), v.strip().strip('"').strip("'")
            if k and k not in os.environ:
                os.environ[k] = v
    except FileNotFoundError:
        pass


try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    _load_env_file(Path(__file__).resolve().parent / ".env")  # stdlib fallback

PARSER_VERSION = "0.3.0"
PROMPT_VERSION = "loop12-v2-betterbriefs"
# A browser-like UA so Cloudflare-fronted APIs (Groq, Cerebras) don't 1010-block us.
_HTTP_UA = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"

# ---------------------------------------------------------------------------
# LLM call ledger — instrumentation only, zero behavioural effect. Reset per
# run(); snapshot lands in brief_object.json → meta.llm_stats so optimisation
# work is measured, not eyeballed. prompt/completion tokens come from provider
# `usage` when present; chars are always counted as the fallback ruler.
# ---------------------------------------------------------------------------
_LLM_STATS = {}
_STATS_LOCK = __import__("threading").Lock()   # loops 3–7 retrieval/rerank run in threads


def _stats_reset():
    with _STATS_LOCK:
        _LLM_STATS.clear()
        _LLM_STATS.update({"calls": 0, "http_attempts": 0, "input_chars": 0,
                           "output_chars": 0, "prompt_tokens": 0, "completion_tokens": 0,
                           "retries": 0, "rate_limited": 0, "by_provider": {}})


def _stats_call(provider_label: str, in_chars: int):
    if not _LLM_STATS:
        _stats_reset()
    with _STATS_LOCK:
        _LLM_STATS["calls"] += 1
        _LLM_STATS["input_chars"] += in_chars
        bp = _LLM_STATS["by_provider"]
        bp[provider_label] = bp.get(provider_label, 0) + 1


def _stats_usage(usage: dict | None, out_chars: int):
    if not _LLM_STATS:
        _stats_reset()
    with _STATS_LOCK:
        _LLM_STATS["output_chars"] += out_chars
        if usage:
            _LLM_STATS["prompt_tokens"] += int(usage.get("prompt_tokens") or 0)
            _LLM_STATS["completion_tokens"] += int(usage.get("completion_tokens") or 0)
HERE = Path(__file__).resolve().parent

# ---------------------------------------------------------------------------
# 1. INGEST
# ---------------------------------------------------------------------------

_VISION_PROMPT = (
    "Transcribe this document image into clean, faithful text for an advertising-brief pipeline.\n"
    "Rules: (1) Capture ALL text verbatim — headings, body, bullets, labels, captions, table cells, "
    "prices, names, figures. Lose nothing. (2) Preserve reading order and structure (use markdown "
    "headings / bullets / tables to mirror the layout). (3) For a meaningful non-text visual (a chart, "
    "an org diagram, a product photo with a caption), add a short bracketed note of what it shows. "
    "(4) Do NOT summarise, interpret, or invent — transcribe only. Output only the transcription."
)

# image input → faithful text, via a NIM vision model (no-loss capture stays intact)
_IMAGE_MIME = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
               ".webp": "image/webp", ".gif": "image/gif", ".bmp": "image/bmp"}


def _vision_transcribe(image_bytes: bytes, mime: str, label: str = "image") -> str:
    """Transcribe one image to text with a vision model (default NIM nemotron-nano-vl).
    Returns '' and warns on failure rather than crashing the run."""
    # Vision endpoint is independent of the main model: point it at NIM (default) or a
    # local Ollama (keyless) via BRIEF_VISION_BASE/MODEL — e.g. gemma3 for image->text.
    base = os.environ.get("BRIEF_VISION_BASE", PROVIDERS["nim"][0])
    model = os.environ.get("BRIEF_VISION_MODEL", "nvidia/llama-3.1-nemotron-nano-vl-8b-v1")
    key = os.environ.get("BRIEF_VISION_API_KEY") or os.environ.get("NVIDIA_API_KEY")
    is_local = "localhost" in base or "127.0.0.1" in base   # e.g. Ollama — no key needed
    if not key and not is_local:
        print(f"[i] {label}: image ingest needs a vision key (NVIDIA_API_KEY / BRIEF_VISION_API_KEY) "
              "or a local endpoint (BRIEF_VISION_BASE=http://localhost:11434/v1) — skipping.",
              file=sys.stderr)
        return ""
    headers = {"Content-Type": "application/json", "User-Agent": _HTTP_UA}
    if key:
        headers["Authorization"] = f"Bearer {key}"
    payload = {
        "model": model, "temperature": 0.0,
        "max_tokens": int(os.environ.get("BRIEF_MAX_TOKENS", "4000")),
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": _VISION_PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{base64.b64encode(image_bytes).decode()}"}},
        ]}],
    }
    req = urllib.request.Request(
        base.rstrip("/") + "/chat/completions", data=json.dumps(payload).encode(),
        headers=headers, method="POST")
    for attempt in range(3):
        try:
            with urllib.request.urlopen(req, timeout=300) as r:
                content = json.loads(r.read())["choices"][0]["message"].get("content") or ""
                return re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", "replace")[:200]
            if e.code in (429, 500, 502, 503) and attempt < 2:
                import time
                time.sleep(5 * (attempt + 1)); continue
            print(f"[i] {label}: vision model HTTP {e.code}: {detail}", file=sys.stderr)
            return ""
        except (urllib.error.URLError, TimeoutError, OSError) as e:
            if attempt < 2:
                import time
                time.sleep(3 * (attempt + 1)); continue
            print(f"[i] {label}: vision model unreachable: {e}", file=sys.stderr)
            return ""
    return ""


def _pdf_vision_transcribe(path: Path, max_pages: int = 20) -> str:
    """Render an image-only / slide-deck PDF page by page and transcribe each."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print(f"[i] {path.name}: looks image-only but PyMuPDF isn't installed "
              "(pip install pymupdf) — can't transcribe.", file=sys.stderr)
        return ""
    doc = fitz.open(str(path))
    n = min(len(doc), max_pages)
    if len(doc) > max_pages:
        print(f"[i] {path.name}: image-PDF — transcribing first {max_pages} of {len(doc)} pages.",
              file=sys.stderr)
    parts = []
    for i in range(n):
        png = doc[i].get_pixmap(dpi=150).tobytes("png")
        t = _vision_transcribe(png, "image/png", f"{path.name} p{i + 1}")
        if t:
            parts.append(f"--- page {i + 1} ---\n{t}")
    return "\n\n".join(parts)


def _strip_html(html: str) -> str:
    """Crude HTML→text: drop tags, unescape the common entities. Good enough for
    an email body when no text/plain part exists."""
    import html as _h
    text = re.sub(r"(?is)<(script|style).*?</\1>", "", html)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p\s*>", "\n\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return _h.unescape(text)


def ingest_email_text(raw: str) -> str:
    """A copy-pasted email is just text. Keep it verbatim — Loop 1 is no-loss —
    but normalise CRLF so the segmenter sees clean lines."""
    return raw.replace("\r\n", "\n").replace("\r", "\n")


def ingest(path: Path) -> tuple[str, str]:
    """Return (raw_text, mime) from .txt/.md, .docx, .pdf, or .eml."""
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".text"}:
        return path.read_text(encoding="utf-8", errors="replace"), "text/plain"
    if suffix == ".eml":
        import email
        from email import policy
        msg = email.message_from_bytes(path.read_bytes(), policy=policy.default)
        body = msg.get_body(preferencelist=("plain", "html"))
        content = body.get_content() if body else (msg.get_content() or "")
        if body is not None and body.get_content_type() == "text/html":
            content = _strip_html(content)
        hdr = [f"{k}: {msg[k]}" for k in ("Subject", "From", "Date") if msg[k]]
        text = ("\n".join(hdr) + "\n\n" + content) if hdr else content
        return ingest_email_text(text), "message/rfc822"
    if suffix == ".docx":
        try:
            import docx
        except ImportError:
            sys.exit("Need python-docx for .docx:  pip install python-docx")
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            sys.exit("Need pdfplumber for .pdf:  pip install pdfplumber")
        out = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                out.append(page.extract_text() or "")
        text = "\n".join(out)
        # Image-only / slide-deck PDFs carry little or no text layer — fall back to
        # rendering each page and transcribing it with the vision model.
        if len(text.strip()) < max(200, 40 * max(1, len(out))):
            print(f"[i] {path.name}: thin text layer — transcribing pages with the vision model.",
                  file=sys.stderr)
            text = _pdf_vision_transcribe(path) or text
        return text, "application/pdf"
    if suffix in _IMAGE_MIME:
        mime = _IMAGE_MIME[suffix]
        return _vision_transcribe(path.read_bytes(), mime, path.name), mime
    sys.exit(f"Unsupported file type: {suffix}. Use .txt, .md, .docx, .pdf, .eml, "
             "or an image (.png/.jpg/.jpeg/.webp) — or paste with --text / '-' for stdin.")


# ---------------------------------------------------------------------------
# 2. SEGMENT
# ---------------------------------------------------------------------------

def segment(text: str) -> list[str]:
    """Coalesce soft-wrapped lines into blocks, then sentence-split. Each
    segment becomes a row in the no-loss ledger."""
    blocks: list[str] = []
    buf: list[str] = []

    def flush():
        if buf:
            blocks.append(" ".join(buf).strip())
            buf.clear()

    for raw in text.splitlines():
        stripped = raw.strip()
        is_bullet = bool(re.match(r"^\s*[-*•]\s+", raw))
        is_label = bool(re.match(r"^[A-Za-z /]{3,30}\s*[:=]\s+\S", stripped))
        if not stripped:
            flush(); continue
        if is_bullet or is_label:
            flush()
        buf.append(stripped.lstrip("-*• \t"))
    flush()

    segs: list[str] = []
    for block in blocks:
        for piece in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\"'])", block):
            piece = piece.strip()
            if len(piece) >= 4:
                segs.append(piece)
    return segs


# ---------------------------------------------------------------------------
# 3a. EXTRACT (heuristic) — Loop 1 faithful capture, no API needed
# ---------------------------------------------------------------------------

LABEL_MAP = {
    "background": "background_context", "context": "background_context",
    "problem": "business_problem", "challenge": "business_problem",
    "objective": "objective", "goal": "objective",
    "audience": "target_audience", "target": "target_audience",
    "budget": "budget", "timeline": "timeline", "timing": "timeline",
    "deadline": "timeline", "deliverable": "deliverables",
    "mandator": "mandatories", "must": "mandatories",
    "kpi": "success_metrics", "success": "success_metrics", "metric": "success_metrics",
    "competitor": "competitors_market", "tone": "tone_and_brand",
    "brand guideline": "tone_and_brand",
}

LIST_FIELDS = {"deliverables", "mandatories", "timeline", "success_metrics",
               "decision_makers", "constraints", "objective", "proof_points",
               "evaluation_criteria"}

KEYWORD_CUES = {
    "business_problem": ["stalled", "perception", "problem", "struggl", "down vs"],
    "target_audience":  ["audience", "switcher", "family", "suburban", "demographic"],
    "budget":           ["£", "$", "€", "budget", "working media"],
    "timeline":         ["week of", "deadline", "pitch presentations", "by july"],
    "mandatories":      ["mandatory", "must ", "can't", "cannot", "asa", "logo",
                         "ci/", "ci ", "claim", "naming", "lockup"],
    "success_metrics":  ["kpi", "consideration", "success =", "test drive", "ipa-style"],
    "decision_makers":  ["cmo", "brand director", "the one to win", "decision-maker"],
    "competitors_market": ["tesla", "polestar", "kia", "hyundai", "competitor", "own \""],
    "tone_and_brand":   ["slogan", "heritage", "unmistakably", "tone of voice"],
}


def _cap(value, status="fact", quote=None, conf=0.6):
    return {"value": value, "status": status, "source_quote": quote, "confidence": conf}


def extract_heuristic(segments):
    fields, used = {}, set()

    def add(field, idx, seg):
        if field in LIST_FIELDS:
            fields.setdefault(field, []).append(_cap(seg, quote=seg))
        elif field not in fields:
            fields[field] = _cap(seg, quote=seg)
        used.add(idx)

    for idx, seg in enumerate(segments):
        low = seg.lower()
        m = re.match(r"^([A-Za-z /]{3,30}?)\s*[:=]\s*(.+)$", seg)
        if m:
            label = m.group(1).strip().lower()
            for key, field in LABEL_MAP.items():
                if key in label:
                    add(field, idx, seg); break
            else:
                pass
            if idx in used:
                continue
        for field, cues in KEYWORD_CUES.items():
            if any(c in low for c in cues):
                add(field, idx, seg); break
    return fields, used


# ---------------------------------------------------------------------------
# 3b. EXTRACT (llm)
# ---------------------------------------------------------------------------

EXTRACTION_SYSTEM = """You are Loop 1 of an ad-agency briefing system, the
Client Brief Parser. Convert a messy client brief into a faithful structured
capture as JSON.

"fields" MUST use ONLY these exact keys (do not invent new field names):
  background_context  - the situation/context behind the brief
  business_problem    - the core problem/challenge to solve (ALWAYS capture this if stated)
  objective           - what the work must achieve   [array, a handful max]
  target_audience     - who we are talking to
  deliverables        - what we must produce        [array]
  mandatories         - non-negotiables: legal, brand, naming, claims  [array]
  budget              - money
  timeline            - dates/deadlines              [array]
  success_metrics     - KPIs / how success is measured  [array]
  key_message         - the ONE single-minded message the client wants to land
  proof_points        - evidence/claims supporting the key message  [array]
  evaluation_criteria - how the client says ideas/work will be judged  [array]
  strategic_angle     - any strategic direction/approach the client suggests
  anti_target         - who the brand is explicitly NOT for / NOT targeting
  competitors_market  - competitors and market context
  tone_and_brand      - tone of voice, brand heritage, style
  decision_makers     - who decides / who to win     [array]
  constraints         - other limits                 [array]
If something does not fit a key, attach it to the CLOSEST key. Never create
keys like "Pitch Details" or "Key Themes". Themes/landmines belong in how_to_win.

Each value is an object: {"value", "status", "source_quote", "confidence"}.
Array fields are lists of such objects.
- status: "fact" (stated), "assumption" (inferred), "gap" (not provided -> value null).
- objective items also carry "objective_type": "commercial" | "behavioural" |
  "attitudinal" (BetterBriefs: the three types should coexist and link —
  attitude shift -> behaviour change -> commercial outcome).
- EVERY fact MUST include a verbatim source_quote copied word-for-word from the brief.
- LOSE NOTHING: every concrete sentence in the brief must be reflected in some
  field's value or source_quote.

how_to_win holds ONLY what the brief reveals (stated_evaluation_criteria,
unstated_needs, likely_landmines, winning_themes, proof_required) as
{"point","evidence"} items. Do NOT invent strategy.

open_questions = what the brief FAILS to answer. Do NOT ask about anything the
brief already states (e.g. if the problem/timeline/decision-maker is given, do
not ask for it).

Return ONLY JSON with keys: fields, how_to_win, open_questions.
Output raw JSON only — no markdown fences, no commentary before or after."""


# BetterBriefs scorecard — judges the CLIENT brief against the BetterBriefs
# rubric (reference/betterbriefs/). Presence is checked elsewhere; this judges
# quality: the dominant failure mode is present-but-VAGUE (78% of marketers
# think their briefs are clear; 5% of agencies agree).
SCORECARD_SYSTEM = """You are a brief-quality judge applying the BetterBriefs
rubric (the global study on briefing) to a CLIENT brief. Judge ONLY what the
brief text says — quote evidence verbatim, do not invent.

Score exactly these dimensions, each as
{"dimension", "verdict": "pass"|"vague"|"missing", "evidence", "fix"}:
  objectives_quality   - a handful at most, benchmarked + time-stamped, the
                         commercial/behavioural/attitudinal chain linked, not
                         wishful, clear hierarchy (not a shopping list)
  audience_vividness   - a vivid picture (demographics + psychographics +
                         needs); FLAG demographic cliches ("millennials",
                         "everyone", bare age ranges) as vague; states who it
                         is NOT for
  single_minded_message - ONE key message, supported by relevant proof points
  evaluation_criteria  - how the work will be judged is stated
  budget_interlock     - budget, objectives and audience are mutually
                         feasible (flag mass-market ambitions on small money)
  strategic_clarity    - a clear strategic angle/choice, including what NOT
                         to do; strategy is not left for the agency to guess
  language             - simple, jargon-free, succinct; no category-speak

Also detect single-mindedness of the WHOLE brief: one brief = one strategy.
If it bundles mutually exclusive strategies or multiple separate jobs
(e.g. several exercises/events/streams), say how to split it.

Return ONLY raw JSON (no fences, no commentary):
{"dimensions": [ ...exactly the 7 above... ],
 "single_mindedness": {"verdict": "single"|"multiple",
                       "split_into": ["one line per separate brief", ...]},
 "summary": "one sentence on overall brief quality"}"""


# The zone-3 strategic fields are GENERATED, not extracted (the guided-generative
# fill). Order follows the schema dependency graph: insight feeds smp, smp feeds
# reasons_to_believe, and desired_response leans on smp + objectives.
GEN_ZONE3_ORDER = ["insight", "smp", "reasons_to_believe", "desired_response"]
# "Do" must be an observable behaviour; these are the schema's own bad_example verbs.
FORBIDDEN_DO_VERBS = ("engage with", "explore the", "interact with", "connect with the brand")


def _field_by_id(schema, fid):
    for f in schema.get("fields", []):
        if f.get("id") == fid:
            return f
    return None


def _gen_field_system(field, n: int = 1) -> str:
    """Build a generation system prompt for ONE golden field straight from the schema
    — no hard-coded sentence template. The schema's good_example carries the shape;
    the bad_example is a hard negative. This is what frees the insight from Mad-Libs.
    n>1 switches to TOURNAMENT mode: one call returns n distinct drafts (the heavy
    context is sent once instead of n times)."""
    mw = field.get("max_words")
    lim = f"Hard limit: {mw} words.\n" if mw else ""
    own = ""
    if field.get("id") in ("insight", "smp"):
        own = ("\nOWNABLE TENSION: claim territory the named competitor does NOT own. If every rival "
               "in the category would nod at your line, it is a category truth and a FAIL — use the "
               "competitor_context to find the white space. Reconcile the WHOLE stated audience (if "
               "it is split, name the tension that unites the segments). Never restate the brand's "
               "own existing vision / mission / tagline; that is not a campaign proposition.\n")
    t = field.get("type")
    if t == "tfd":
        out_shape = ('{"think": "...", "feel": "...", '
                     '"do": "<a concrete, observable behaviour — NOT \'engage\'/\'explore\'>"}')
    elif t == "list":
        out_shape = '["...", "...", "..."]'
    else:
        out_shape = '"<text>"'
    return (
        f"You are a senior strategy planner writing the '{field['label']}' field of a brief "
        f"for THIS specific brand. Write it now — do not extract it, derive it.\n\n"
        f"WHAT THIS FIELD IS: {field.get('prompt','')}\n{lim}\n"
        f"STYLE REFERENCE (a DIFFERENT brand — copy the depth/shape ONLY, never its words, "
        f"brand, topic; and NEVER mention it in your rationale):\n"
        f"  {field.get('good_example','')}\n\n"
        f"BAD — never produce anything like this:\n"
        f"  {field.get('bad_example','')}  ({field.get('bad_reason','')})\n"
        f"{own}\n"
        "Reason from the brief context and the real award-winning PRECEDENTS provided below. The "
        "precedents are for SHAPE and DEPTH only — do NOT borrow their words, brands or themes. Be "
        "specific to this brand: a line that could be pasted onto a different brief is a failure. Do "
        "NOT output a generic fill-in-the-blank sentence. You are synthesising strategy (source "
        "'inferred'), never inventing client facts.\n\n"
        + (f'Return ONLY raw JSON, no fences: {{"value": {out_shape}, '
           '"confidence": 0.0-1.0, "rationale": "one line on THIS brand\'s tension and which AWARD '
           'PRECEDENT shaped it — never mention the style reference"}'
           if n <= 1 else
           f"You will write {n} GENUINELY DISTINCT drafts of this field — different strategic "
           f"ideas, not rewordings of one idea. Make them compete.\n"
           f'Return ONLY raw JSON, no fences: {{"candidates": [{n} objects, each '
           f'{{"value": {out_shape}, "confidence": 0.0-1.0, "rationale": "one line — never '
           f'mention the style reference"}}]}}')
    )


def _coerce_candidates(o) -> list:
    """Accept every plausible shape a batched tournament call can come back in:
    {"candidates":[...]}, a bare list, or a single {"value": ...} draft."""
    if isinstance(o, dict) and isinstance(o.get("candidates"), list):
        return [c for c in o["candidates"] if isinstance(c, dict) and c.get("value")]
    if isinstance(o, list):
        return [c for c in o if isinstance(c, dict) and c.get("value")]
    if isinstance(o, dict) and o.get("value"):
        return [o]
    return []


def _build_golden_system():
    """Build a per-field extraction system prompt from golden_brief.schema.json."""
    schema_path = HERE / "golden-brief" / "golden_brief.schema.json"
    if not schema_path.exists():
        return None
    schema = json.loads(schema_path.read_text())
    lines = [
        "You are a senior strategic planner filling a Golden Brief from a raw client brief.",
        "Extract ONLY what the brief contains. Mark provenance accurately.",
        "",
        "PROVENANCE:",
        "  client_stated — directly from brief text; include verbatim source_quote",
        "  inferred      — reasonably implied but not stated",
        "  missing       — not in brief at all; value MUST be null",
        "",
        "CONFIDENCE: 0.9+ verbatim, 0.6–0.8 inferred, 0.0 missing.",
        "",
        "FIELD SPECIFICATIONS (fill all 11 content fields):",
    ]
    for f in schema.get("fields", []):
        t = f.get("type", "text")
        if t == "objectives":
            shape = '{"commercial": "str", "behavioural": "str", "attitudinal": "str"}'
        elif t == "list":
            shape = '["item1", "item2", ...]'
        elif t == "tfd":
            shape = '{"think": "str", "feel": "str", "do": "str"}'
        else:
            shape = "string or null"
        lines += [
            f"\n## {f['id']}  ({f['label']})",
            f"Instruction: {f['prompt']}",
            f"Good: {f['good_example']}",
            f"Bad: {f['bad_example']} — Why bad: {f['bad_reason']}",
            f"Max words: {f.get('max_words', 'no limit')}  Value shape: {shape}",
        ]
    lines += [
        "",
        "NOTE: insight, smp, reasons_to_believe, desired_response are zone-3 STRATEGY",
        "fields. Mark them client_stated ONLY if the brief explicitly articulates that",
        "strategic element in its own words. A market fact, background, objective, or",
        "audience description is NOT an insight or a proposition — if the brief merely",
        "describes the situation, mark these fields MISSING (value null). Never repackage",
        "a background/market statement as the insight or SMP. Do NOT invent strategy.",
        "",
        'Return ONLY raw JSON, no fences:',
        '{"fields": {"<id>": {"value": <value>, "source": "client_stated"|"inferred"|"missing",',
        '  "confidence": 0.0-1.0, "source_quote": "verbatim" | null}, ...}}',
    ]
    return "\n".join(lines)


def extract_golden_brief(raw_text: str) -> "dict | None":
    """Per-field Golden Brief extraction using schema prompts + good/bad examples.
    Returns a dict with 'fields' key, or None on failure."""
    import time
    system = _build_golden_system()
    if not system:
        return None
    user = f"CLIENT BRIEF:\n\"\"\"\n{_clip_brief(raw_text)}\n\"\"\""
    # 3 retries with short backoff — model sometimes returns inconsistent output
    # when called in rapid succession during a full pipeline run
    for attempt in range(3):
        if attempt:
            time.sleep(2)
        obj = _json_call(user, system=system, retries=0, max_tokens=MAXTOK_EXTRACT)
        if isinstance(obj, dict) and isinstance(obj.get("fields"), dict):
            return obj
    return None


def _fv(field) -> str:
    """Safely extract .value from a golden field that may be a dict or plain string."""
    if isinstance(field, dict):
        return str(field.get("value") or "")
    return str(field or "")


def _word_count(v) -> int:
    if isinstance(v, dict):
        return sum(_word_count(x) for x in v.values())
    if isinstance(v, list):
        return sum(_word_count(x) for x in v)
    return len(str(v).split())


def _text_overlap(a: str, b: str) -> float:
    """Fraction of a's (3+ char) words that also appear in b. Used to detect a
    strategy field that just parrots a client fact."""
    wa = {w for w in re.findall(r"[a-z]{3,}", a.lower())}
    if not wa:
        return 0.0
    wb = {w for w in re.findall(r"[a-z]{3,}", b.lower())}
    return len(wa & wb) / len(wa)


_BRAND_BOILERPLATE_MARKERS = ("vision", "mission", "purpose", "brand promise",
                              "brand value", "campaign claim", "brand claim", "tagline")


def _brand_boilerplate(text: str) -> str:
    """Collect the brand's own vision / mission / claim / tagline lines from the brief
    (incl. attachments). Used to reject an SMP that just echoes the masterbrand line —
    an SMP must be a campaign CHOICE, not a restatement of the standing brand vision."""
    lines = []
    for ln in (text or "").splitlines():
        low = ln.lower()
        if any(m in low for m in _BRAND_BOILERPLATE_MARKERS):
            lines.append(ln.strip())
    return " ".join(lines)


def _rubric_gate(field, value, brand_lines: str = "", ctx: str = "") -> "tuple[bool, list]":
    """Run a field's own schema rubric against a generated value. Auto tests run
    inline; the llm tests run in one judge call. Returns (passed, failures).
    A hard failure (over limit / forbidden verb / masterbrand echo) fails outright;
    up to one soft (llm) failure is tolerated so a subjective judge can't nuke every field.
    `ctx` (the field's upstream deps, e.g. the insight + competitor_context) is given to
    the judge so derivation/ownability tests are checked against real context, not blind."""
    rubric = field.get("rubric") or []
    mw = field.get("max_words")
    blob = (json.dumps(value).lower() if not isinstance(value, str) else value.lower())
    hard, soft = [], []
    if mw and _word_count(value) > mw * 1.2:
        hard.append(f"over the {mw}-word limit ({_word_count(value)} words)")
    if field.get("type") == "tfd" and any(v in blob for v in FORBIDDEN_DO_VERBS):
        hard.append("'do' is not an observable behaviour (engage/explore/interact)")
    # The SMP must be a campaign choice, not a restatement of the masterbrand vision.
    if field.get("id") == "smp" and brand_lines and isinstance(value, str) \
            and _text_overlap(value, brand_lines) >= 0.5:
        hard.append("SMP echoes the masterbrand vision/claim — needs a campaign-specific proposition")
    llm_tests = [r for r in rubric if r.get("method") == "llm"]
    if llm_tests:
        tests = "\n".join(f'- {r["id"]}: {r["test"]}' for r in llm_tests)
        judge = _json_call(
            (f"UPSTREAM CONTEXT (use this to judge derivation/ownability — do NOT re-test it):\n{ctx}\n\n"
             if ctx else "")
            + f"FIELD VALUE:\n{json.dumps(value)}\n\nTESTS:\n{tests}",
            system=("You are a fair but rigorous brief-quality judge. Judge the VALUE on each test, using "
                    "the upstream context where given (e.g. verify a proposition derives from the stated "
                    "insight and is ownable against the stated competitor — do not fail derivation merely "
                    "because the context wasn't repeated in the line). For each test decide pass/fail. "
                    'Return ONLY raw JSON: {"<test_id>": {"pass": true|false, "why": "short"}}'),
            retries=1, max_tokens=MAXTOK_JUDGE)
        if isinstance(judge, dict):
            for r in llm_tests:
                res = judge.get(r["id"]) or {}
                if isinstance(res, dict) and res.get("pass") is False:
                    soft.append(f'{r["id"]}: {res.get("why", "failed")}')
    passed = (not hard) and len(soft) < 2
    return passed, hard + soft


def _judge_hero_candidates(field, candidates, brand_lines: str = ""):
    """Rank candidate values for a hero field (insight / smp) by PURITY and ownability
    with one LLM judge call, and return them reordered best-first. Best-effort: on any
    failure the candidates are returned unchanged (the rubric gate is still authoritative)."""
    if len(candidates) < 2:
        return candidates
    llm_tests = [r for r in (field.get("rubric") or []) if r.get("method") == "llm"]
    crit = "; ".join(f'{r["id"]}: {r["test"]}' for r in llm_tests) \
        or "single-minded (one idea, not a list); ownable vs the competitor; specific to this brand"
    listing = "\n".join(f"[{i}] {json.dumps(c.get('value'))}" for i, c in enumerate(candidates))
    judge = _json_call(
        f"FIELD: {field['label']}\nGOOD shape (different brand, do not copy): {field.get('good_example','')}\n"
        f"BAD: {field.get('bad_example','')} ({field.get('bad_reason','')})\nJUDGE EACH ON: {crit}\n\n"
        f"CANDIDATES:\n{listing}",
        system=("You are a strategy director ranking candidate '" + field["label"] + "' lines for a "
                "creative brief. Reward PURITY and single-mindedness (one idea, never a list or an "
                "'and'), ownable territory (a direct rival could not say the same line), and a real "
                "human tension specific to THIS brand. Penalise category truths everyone would nod "
                "at, restated brand taglines, and anything trying to say two things. Return ONLY raw "
                'JSON: {"ranking": [candidate indexes, best first], "why": "one line on the winner"}'),
        retries=1, max_tokens=MAXTOK_JUDGE)
    if isinstance(judge, dict) and isinstance(judge.get("ranking"), list):
        order = [i for i in judge["ranking"] if isinstance(i, int) and 0 <= i < len(candidates)]
        order += [i for i in range(len(candidates)) if i not in order]   # append any the judge dropped
        ranked = [candidates[i] for i in order]
        if ranked and isinstance(judge.get("why"), str):
            ranked[0] = {**ranked[0], "_judge_why": judge["why"]}
        return ranked
    return candidates


def _refine_field(field, value, note: str = ""):
    """One sharpening pass on the chosen hero value — purer, more single-minded, more
    ownable. Returns a candidate dict, or None on failure (caller keeps the original)."""
    system = _gen_field_system(field) + (
        "\n\nREFINE MODE: you are given a strong draft. Make it PURER and more single-minded — "
        "one idea only, sharper, more ownable. Keep what already works; never add a second idea. "
        "If it is already optimal, return it unchanged.")
    user = (f"DRAFT '{field['label']}' to sharpen:\n{json.dumps(value)}\n"
            + (f"\nDirector's note: {note}\n" if note else "")
            + "Return the improved value in the same JSON shape.")
    raw = _json_call(user, system=system, max_tokens=MAXTOK_GEN)
    return raw if isinstance(raw, dict) and raw.get("value") else None


# Distinct strategic facets seeded one-per-draft so the SMP tournament gets a real
# spread of ideas instead of N identical draws (decision / perception / ambition / craft).
SMP_ANGLE_SEEDS = (
    "lead with the contrast between what this audience settles for and what they could deliberately choose",
    "lead with how this audience is judged or perceived by others — the product as a verdict on their standards",
    "lead with the competitive edge or advantage the product gives them over rivals",
    "lead with the brand's own craft / design / engineering equity as the reason to choose it",
)


def _smp_territory(brief_text: str, competitor_ctx: str) -> dict:
    """Map the SMP's ownable white space in one call. Returns {own, avoid, rival}:
    `own`  — the territory THIS brand should claim (its white space, per the brief);
    `avoid`— the emotional/territorial ground the named competitor ALREADY owns;
    `rival`— the competitor's name (for the 'could they run this line?' kill-test).
    Best-effort: falls back to a generic ownable framing if the call fails."""
    fallback = {"own": "the deliberate, considered choice a direct rival cannot credibly claim",
                "avoid": (competitor_ctx or "the category's generic, everyone-says-it territory"),
                "rival": "the named competitor"}
    if not resolve_provider() or not (brief_text or competitor_ctx):
        return fallback
    obj = _json_call(
        f"BRIEF:\n\"\"\"\n{_clip_brief(brief_text or '')}\n\"\"\"\n\nCOMPETITOR CONTEXT: {competitor_ctx}",
        system=("You map strategic white space for a single-minded proposition. From the brief and the "
                "competitor context identify three things: the named competitor; the emotional/territorial "
                "ground that competitor ALREADY OWNS (so we steer away from it — give the concept plus its "
                "common synonyms); and the adjacent white space THIS brand should claim instead (its real, "
                "ownable edge as the brief itself describes it). Be concrete and short. Return ONLY raw "
                'JSON: {"rival": "competitor name", "avoid": "the concept they own + synonyms", '
                '"own": "the white space this brand should claim"}'),
        retries=1, max_tokens=MAXTOK_GEN)
    if isinstance(obj, dict) and obj.get("own") and obj.get("avoid"):
        return {"own": str(obj["own"]), "avoid": str(obj["avoid"]),
                "rival": str(obj.get("rival") or fallback["rival"])}
    return fallback


def _smp_territory_gate(value, territory: dict) -> "tuple[bool, str]":
    """Semantic kill-test for the SMP (a lexical word-ban misses synonyms — 'never lets
    you down' is the rival's ground without the banned words). One judge call runs the two
    CD-mandated tests: a POSITIVE gate (does it live on the brand's white space?) and a
    COMPETITOR test (could the rival run this line verbatim?). Returns (passed, reason).
    Best-effort: a judge-call failure does NOT block (the rubric gate still applies)."""
    if not isinstance(value, str) or not value.strip():
        return False, "empty"
    obj = _json_call(
        f"PROPOSITION: {value}\nBRAND SHOULD OWN: {territory['own']}\n"
        f"{territory['rival']} ALREADY OWNS: {territory['avoid']}",
        system=("You are a strategy director enforcing ownable territory for a proposition. Run two tests.\n"
                "own_territory: does the line clearly live on what the BRAND should own, rather than on the "
                "competitor's ground — counting synonyms and rephrasings, not just exact words?\n"
                "competitor_could_run: would the line FAIL because the named competitor's own campaign "
                "could run it verbatim without changing its meaning?\n"
                'Return ONLY raw JSON: {"own_territory": true|false, "competitor_could_run": true|false, '
                '"why": "one line"}'),
        retries=1, max_tokens=MAXTOK_JUDGE)
    if not isinstance(obj, dict):
        return True, ""   # judge unavailable — don't block on infra failure
    on = (obj.get("own_territory") is not False) and (obj.get("competitor_could_run") is not True)
    return on, str(obj.get("why") or ("walks onto the competitor's ground" if not on else ""))


def _precedent_blocks(loops: dict, key: str):
    """Pull a loop's retrieved evidence into (ipa_block, method_block, evidence_ids):
    IPA effectiveness cases (shape/depth exemplars) and playbook/framework snippets.
    Each hero field reads ITS OWN loop — insight←loop4 (tension), smp←loop5 (the
    proposition playbook, incl. the single-minded-proposition rulebook)."""
    loop = (loops or {}).get(key) or {}
    ipa_ex, methods, ev_ids = [], [], []
    for e in (loop.get("evidence") or []):
        snip = (e.get("snippet") or "").strip()
        if not snip:
            continue
        src = e.get("source") or e.get("framework") or e.get("citation") or ""
        if (e.get("category") or "") == "ipa_effectiveness_case":
            ipa_ex.append(f"- {snip[:220]}")
            if src:
                ev_ids.append(src)
        else:
            methods.append(f"- {snip[:160]}")
    return ("\n".join(ipa_ex[:5]) or "(no IPA precedent retrieved)",
            "\n".join(methods[:3]) or "(no playbook evidence)", ev_ids)


def fill_derivable_fields(golden_fields: dict, loop37_result: dict, schema: dict, brief_text: str = ""):
    """Guided-generative fill of the zone-3 strategy fields (insight → smp →
    reasons_to_believe → desired_response), schema-driven via each field's
    depends_on and rubric. A field is generated ONLY if its extracted source is
    missing/inferred — a client_stated value is never overwritten (the no-invent
    invariant). Each generated value is rubric-gated; a failure downgrades the field
    to 'missing' and surfaces an open question. Mutates golden_fields in place so a
    later field can read a freshly generated upstream one (smp reads insight).
    Returns (fills, open_questions)."""
    if not resolve_provider():
        return {}, []
    floor = float(schema.get("confidence_floor") or 0.6)

    # Precedent pools: each hero field draws on ITS OWN loop's retrieval. The insight reads
    # Loop 4 (human tension); the SMP reads Loop 5 (the proposition playbook — incl. the
    # single-minded-proposition rulebook). SMP falls back to Loop 4 if Loop 5 was empty.
    loops = loop37_result.get("loops") or {}
    insight_ipa, insight_methods, insight_ev = _precedent_blocks(loops, "loop4_insight")
    smp_ipa, smp_methods, smp_ev = _precedent_blocks(loops, "loop5_proposition")
    if smp_ipa.startswith("(no") and smp_methods.startswith("(no"):
        smp_ipa, smp_methods, smp_ev = insight_ipa, insight_methods, insight_ev

    def val(fid):
        f = golden_fields.get(fid)
        return _fv(f) if f else ""

    # Guard: the extractor sometimes mislabels a background/market fact as a strategy
    # field (client_stated). If a strategy value substantially echoes the client facts,
    # downgrade it to missing so generation owns it — nothing is lost, the text still
    # lives in its real fact field.
    fact_blob = " ".join(val(f) for f in
                         ("background", "audience", "objectives", "competitor_context"))
    brand_blob = _brand_boilerplate(brief_text)
    for fid in GEN_ZONE3_ORDER:
        cur = golden_fields.get(fid) or {}
        if isinstance(cur, dict) and cur.get("source") == "client_stated":
            v = _fv(cur)
            if _text_overlap(v, fact_blob) >= 0.6:
                golden_fields[fid] = {
                    "value": None, "source": "missing",
                    "reason": "extracted value echoed a client fact, not a distinct strategy statement",
                }
            elif fid in ("smp", "insight") and brand_blob and _text_overlap(v, brand_blob) >= 0.5:
                # The masterbrand vision/tagline is not a campaign proposition — force generation.
                golden_fields[fid] = {
                    "value": None, "source": "missing",
                    "reason": "echoed the masterbrand vision/claim, not a campaign-specific proposition",
                }

    fills, open_qs = {}, []
    for fid in GEN_ZONE3_ORDER:
        field = _field_by_id(schema, fid)
        if not field:
            continue
        cur = golden_fields.get(fid) or {}
        if isinstance(cur, dict) and cur.get("source") == "client_stated":
            continue  # never overwrite a client fact

        deps = list(field.get("depends_on", []))
        # The sharpest captured thinking (the white space vs the competitor) must reach
        # the insight/SMP generator, not sit in a fact field it never reads.
        if fid in ("insight", "smp") and "competitor_context" not in deps:
            deps.append("competitor_context")
        ctx = "\n".join(f"{d}: {val(d)}" for d in deps if val(d))
        ctx = (ctx + f"\nbackground: {val('background')}").strip()
        use_ipa = fid in ("insight", "smp")
        if fid == "smp":
            f_ipa, f_methods, f_ev = smp_ipa, smp_methods, smp_ev
            rules_label = "PROPOSITION RULEBOOK (apply these rules — how a single-minded proposition is written)"
        else:
            f_ipa, f_methods, f_ev = insight_ipa, insight_methods, insight_ev
            rules_label = "PLANNING FRAMEWORKS"
        user = (
            "BRIEF CONTEXT:\n" + ctx + "\n\n"
            + (f"AWARD-WINNING PRECEDENT (shape & depth only — do not copy):\n{f_ipa}\n\n"
               f"{rules_label}:\n{f_methods}\n\n" if use_ipa else "")
            + f"Write the '{field['label']}' for THIS brand now."
        )
        system = _gen_field_system(field)

        # Hero fields (insight, smp) run a TOURNAMENT: generate N candidates, rank them by
        # purity/ownability, pick the best that clears the rubric, then one sharpen pass.
        # Non-hero fields generate once. N is tunable (a stronger model needs fewer).
        n_cand = int(os.environ.get("BRIEF_HERO_CANDIDATES", "4")) if fid in ("insight", "smp") else 1
        # The SMP is the brief's hardest field to own — map its white space once, generate
        # a wider, angle-seeded spread, and gate every candidate on territory (below).
        territory = None
        if fid == "smp":
            n_cand = max(n_cand, int(os.environ.get("BRIEF_SMP_CANDIDATES", "6")))
            territory = _smp_territory(brief_text, val("competitor_context"))
        # SMP territory block — built once, reused by the batched call and the fallback.
        terr_block = ""
        if fid == "smp" and territory:
            terr_block = (
                f"\n\nOWNABLE TERRITORY — CLAIM THIS: {territory['own']}.\n"
                f"DO NOT walk onto {territory['rival']}'s ground ({territory['avoid']}); a "
                f"proposition {territory['rival']} could also run is a FAIL — find the white space.\n"
                f"This is a single-minded PROPOSITION — the ONE thing to make the audience believe — "
                f"and it MUST visibly derive from the insight above (the reader should see the line "
                f"through to the insight). Write the strategic proposition itself, NOT finished ad "
                f"copy, a headline or a tagline.")

        candidates = []
        if n_cand > 1:
            # BATCHED tournament: ONE call returns all N drafts, so the heavy context
            # (brief ctx + precedent + rulebook) is sent once instead of N times (~80%
            # input cut on the heaviest phase) — and the model can differentiate its own
            # drafts, which gives a wider spread than N independent samples.
            u = user + terr_block
            if fid == "smp" and territory:
                seeds = "\n".join(
                    f"  draft {i+1} — pull the idea this way (do not name the angle): "
                    f"{SMP_ANGLE_SEEDS[i % len(SMP_ANGLE_SEEDS)]}" for i in range(n_cand))
                u += f"\nANGLES — one per draft:\n{seeds}"
            batched = _json_call(u, system=_gen_field_system(field, n=n_cand),
                                 max_tokens=MAXTOK_GEN * 2,
                                 accept=lambda o: bool(_coerce_candidates(o)))
            candidates = _coerce_candidates(batched)[:n_cand]
        if not candidates:
            # Single-draft path: non-hero fields, or fallback when the batched call
            # exhausted the chain (never fail the tournament over one bad response).
            for i in range(max(1, n_cand)):
                u = user + terr_block
                if fid == "smp" and territory:
                    seed = SMP_ANGLE_SEEDS[i % len(SMP_ANGLE_SEEDS)]
                    u += f"\nANGLE FOR THIS DRAFT (pull the idea this way, do not name the angle): {seed}"
                raw = _json_call(u, system=system, max_tokens=MAXTOK_GEN)
                if isinstance(raw, dict) and raw.get("value"):
                    candidates.append(raw)
        if not candidates:
            golden_fields[fid] = {"value": None, "source": "missing",
                                  "reason": "generation produced no output"}
            open_qs.append({"question": f"Agree the {field['label'].lower()} — none could be derived.",
                            "priority": "high" if field.get("hero") else "medium", "blocks_field": fid})
            continue

        if fid in ("insight", "smp"):
            candidates = _judge_hero_candidates(field, candidates, brand_blob)   # reorder best-first

        chosen, chosen_fail = None, None
        for c in candidates:
            ok, fails = _rubric_gate(field, c["value"], brand_lines=brand_blob, ctx=ctx)
            if ok and fid == "smp" and territory:
                on_terr, why = _smp_territory_gate(c["value"], territory)
                if not on_terr:
                    ok, fails = False, [f"walks onto the competitor's ground: {why}"]
            if ok:
                chosen, chosen_fail = c, []
                break
            if chosen is None:
                chosen, chosen_fail = c, fails
        # SMP territory rescue: if no candidate could both pass the rubric AND hold the white
        # space, push the best draft off the competitor's ground once before giving up.
        if fid == "smp" and territory and chosen_fail:
            note = (f"This proposition walks onto {territory['rival']}'s ground ({territory['avoid']}). "
                    f"Rewrite it to claim the brand's own white space: {territory['own']}. Keep the same "
                    f"underlying insight, ONE idea only, within the word limit — a line "
                    f"{territory['rival']} could not credibly run. Make it a strategic PROPOSITION that "
                    f"derives from the insight, not a finished tagline or headline.")
            resc = _refine_field(field, chosen["value"], note=note)
            if resc:
                ok3, _f3 = _rubric_gate(field, resc["value"], brand_lines=brand_blob, ctx=ctx)
                if ok3 and _smp_territory_gate(resc["value"], territory)[0]:
                    chosen = {**resc, "_judge_why": chosen.get("_judge_why", "")}
                    chosen_fail = []
        conf = float(chosen.get("confidence") or floor)

        if chosen_fail or conf < floor:
            golden_fields[fid] = {
                "value": None, "source": "missing",
                "reason": "; ".join(chosen_fail) or f"confidence {conf:.2f} < floor {floor}",
                "rejected_attempt": chosen.get("value"),
            }
            open_qs.append({"question": f"Agree the {field['label'].lower()}.",
                            "why_it_matters": "; ".join(chosen_fail) or "below confidence floor",
                            "priority": "high" if field.get("hero") else "medium", "blocks_field": fid})
            continue

        # Sharpen the winning hero line once; keep the refinement only if it still clears the gate.
        if fid in ("insight", "smp"):
            refined = _refine_field(field, chosen["value"], chosen.get("_judge_why", ""))
            if refined:
                ok2, _f2 = _rubric_gate(field, refined["value"], brand_lines=brand_blob, ctx=ctx)
                if ok2 and float(refined.get("confidence") or conf) >= floor:
                    # Never let the sharpen pass drift the SMP back onto the competitor's ground.
                    keep = True if fid != "smp" or not territory \
                        else _smp_territory_gate(refined["value"], territory)[0]
                    if keep:
                        refined["_judge_why"] = chosen.get("_judge_why", "")
                        chosen, conf = refined, float(refined.get("confidence") or conf)

        entry = {"value": chosen["value"], "source": "inferred",
                 "method": f"gen:{fid}", "confidence": round(conf, 2)}
        if chosen.get("rationale"):
            entry["rationale"] = chosen["rationale"]
        if chosen.get("_judge_why"):
            entry["judge_note"] = chosen["_judge_why"]
        if use_ipa and f_ev:
            entry["evidence_ids"] = f_ev[:5]
        alts = [c["value"] for c in candidates if c.get("value") != chosen["value"]]
        if alts:
            entry["alternatives"] = alts[:3]
        fills[fid] = entry
        golden_fields[fid] = entry  # downstream deps see the generated value
    return fills, open_qs


# provider -> (default base_url, api-key env var, default model)
# nim = NVIDIA NIM (Nemotron) via build.nvidia.com — free for prototyping on Inception.
PROVIDERS = {
    "nim":    ("https://integrate.api.nvidia.com/v1", "NVIDIA_API_KEY",
               "nvidia/llama-3.1-nemotron-70b-instruct"),
    # Free-tier, OpenAI-compatible alternatives — faster/steadier than NIM for this
    # workload. Get a free key, set BRIEF_PROVIDER + the key env var, done.
    "groq":     ("https://api.groq.com/openai/v1", "GROQ_API_KEY", "llama-3.3-70b-versatile"),
    "cerebras": ("https://api.cerebras.ai/v1", "CEREBRAS_API_KEY", "llama-3.3-70b"),
    "gemini":   ("https://generativelanguage.googleapis.com/v1beta/openai",
                 "GEMINI_API_KEY", "gemini-2.0-flash"),
    "openai": ("https://api.openai.com/v1", "OPENAI_API_KEY", "gpt-4o"),
    "ollama": ("http://localhost:11434/v1", None, "llama3.1"),
}


def resolve_provider() -> str:
    """Explicit BRIEF_PROVIDER wins; otherwise auto-detect from whichever key is set."""
    p = os.environ.get("BRIEF_PROVIDER", "").lower().strip()
    if p:
        return p
    for env, prov in (("GROQ_API_KEY", "groq"), ("CEREBRAS_API_KEY", "cerebras"),
                      ("GEMINI_API_KEY", "gemini"), ("NVIDIA_API_KEY", "nim"),
                      ("ANTHROPIC_API_KEY", "anthropic"), ("OPENAI_API_KEY", "openai")):
        if os.environ.get(env):
            return prov
    return ""


def model_for(provider: str) -> str:
    if os.environ.get("BRIEF_MODEL"):
        return os.environ["BRIEF_MODEL"]
    if provider == "anthropic":
        return "claude-opus-4-6"
    return PROVIDERS.get(provider, (None, None, "?"))[2]


# One consistent clip for every prompt that embeds the brief. Judge/strategy calls
# (scorecard, golden extract, territory) don't need the tail of a long deck; extraction
# gets a much more generous window because Loop 1 is the no-loss capture (the project's
# one hard rule) — starving it would drop segments from the ledger's LLM mapping.
CLIP_JUDGE = int(os.environ.get("BRIEF_CLIP_CHARS", "6500"))
CLIP_EXTRACT = int(os.environ.get("BRIEF_CLIP_EXTRACT_CHARS", "12000"))


def _clip_brief(text: str, limit: int = CLIP_JUDGE) -> str:
    """Clip at a sentence/line boundary near the limit so the model never sees a
    mid-word truncation."""
    if not text or len(text) <= limit:
        return text
    cut = text[:limit]
    for m in (cut.rfind(". "), cut.rfind(".\n"), cut.rfind("\n")):
        if m > limit * 0.85:
            return cut[:m + 1]
    return cut


def _user_msg(text, schema):
    return (f"SCHEMA KEYS:\n{json.dumps(schema, indent=2)[:2500]}\n\n"
            f"CLIENT BRIEF:\n\"\"\"\n{_clip_brief(text, CLIP_EXTRACT)}\n\"\"\"")


class _RateLimited(RuntimeError):
    """A link returned HTTP 429. Raised fast (no backoff) so the chain fails over and the
    link is put on a short cooldown — see _LINK_COOLDOWN."""


# Right-sized output budgets per call class. max_tokens counts against free-tier
# TPM budgets (Groq bills the CAP, not actual output), so a global 4000 was ~60%
# waste — judges return ~100-token verdicts. BRIEF_MAX_TOKENS overrides everything.
MAXTOK_JUDGE = 500       # rubric gates / judges / territory gate / rerank: tiny JSON verdicts
MAXTOK_GEN = 1200        # candidate generation / refine: one field's worth of text
MAXTOK_EXTRACT = 2500    # extraction / scorecard / golden / loop synthesis: big JSON


def _chat_openai_compatible(base_url, key, model, user, provider_label="llm",
                            timeout=300, system=None, max_tokens=None, json_mode=False):
    """One code path for NVIDIA NIM, OpenAI, and Ollama — all OpenAI-compatible."""
    # NOTE: do NOT prepend a "detailed thinking off" system message for
    # Nemotron — on NIM a second system message displaces the real one and
    # the model ignores the JSON instruction entirely (verified 2026-06-10).
    messages = [{"role": "system", "content": system or EXTRACTION_SYSTEM},
                {"role": "user", "content": user}]
    payload = {
        "model": model, "temperature": 0.2,
        "max_tokens": int(os.environ.get("BRIEF_MAX_TOKENS") or max_tokens or MAXTOK_EXTRACT),
        "messages": messages,
    }
    # Structured-output mode: cuts the unclean-JSON retries that multiply calls through
    # the chain. Only for providers known to accept it; a 400 retries without it below.
    if json_mode and provider_label.split(":", 1)[0] in ("groq", "cerebras", "openai"):
        payload["response_format"] = {"type": "json_object"}
    # Disable thinking only for reasoning models — non-reasoning models (e.g.
    # llama-3.3-70b) don't support this flag and may error on it.
    if "reasoning" in model or "thinking" in model:
        payload["chat_template_kwargs"] = {"enable_thinking": False}
    body = json.dumps(payload).encode()
    # Groq/Cerebras sit behind Cloudflare, which blocks Python's default UA (error 1010).
    headers = {"Content-Type": "application/json", "User-Agent": _HTTP_UA}
    if key:
        headers["Authorization"] = f"Bearer {key}"
    req = urllib.request.Request(base_url.rstrip("/") + "/chat/completions",
                                 data=body, headers=headers, method="POST")
    _stats_call(provider_label, len(system or EXTRACTION_SYSTEM) + len(user))
    # Retry transient timeouts/network blips — these (not API errors) are what was
    # silently dropping briefs to heuristic mode on NIM.
    for attempt in range(3):
        with _STATS_LOCK:
            _LLM_STATS["http_attempts"] += 1
            if attempt:
                _LLM_STATS["retries"] += 1
        try:
            with urllib.request.urlopen(req, timeout=timeout) as r:
                data = json.loads(r.read())
                msg = data["choices"][0]["message"]
                content = msg.get("content") or ""
                # Reasoning models (e.g. nemotron-*-reasoning) emit <think>…</think>
                # before the answer. Strip it so _loads_lenient finds clean JSON.
                content = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
                _stats_usage(data.get("usage"), len(content))
                return content
        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", "replace")[:400]
            # 429 = rate-limited: fail FAST so the chain fails over to a free sibling link
            # immediately, instead of burning ~15s backing off a provider that's saturated.
            if e.code == 429:
                with _STATS_LOCK:
                    _LLM_STATS["rate_limited"] += 1
                raise _RateLimited(f"HTTP 429 from {provider_label}") from None
            # A 400 right after adding response_format = this model rejects structured
            # output — drop the flag and retry the same link (don't burn a chain hop).
            if e.code == 400 and "response_format" in payload and attempt < 2:
                payload.pop("response_format", None)
                body = json.dumps(payload).encode()
                req = urllib.request.Request(base_url.rstrip("/") + "/chat/completions",
                                             data=body, headers=headers, method="POST")
                continue
            # Transient server errors: one short backoff, then hand off to the chain.
            if e.code in (500, 502, 503) and attempt < 2:
                import time
                print(f"[i] {provider_label} HTTP {e.code} (attempt {attempt+1}/3), backing off…",
                      file=sys.stderr)
                time.sleep(2 * (attempt + 1))
                continue
            raise RuntimeError(f"HTTP {e.code} from {provider_label}: {detail}") from None
        except (urllib.error.URLError, TimeoutError, OSError) as e:
            if attempt < 2:
                import time
                print(f"[i] {provider_label} timeout/blip (attempt {attempt+1}/3), retrying…",
                      file=sys.stderr)
                time.sleep(2 * (attempt + 1))
                continue
            raise RuntimeError(f"{provider_label} unreachable after 3 attempts: {e}") from None


def list_models(provider="nim"):
    """Connectivity check: list available models (esp. Nemotron). For `--check`."""
    default_base, key_env, _ = PROVIDERS[provider]
    base_url = os.environ.get("BRIEF_BASE_URL", default_base)
    key = os.environ.get(key_env) if key_env else "ollama"
    req = urllib.request.Request(base_url.rstrip("/") + "/models",
                                 headers={"Authorization": f"Bearer {key}", "User-Agent": _HTTP_UA})
    with urllib.request.urlopen(req, timeout=30) as r:
        ids = [m["id"] for m in json.loads(r.read()).get("data", [])]
    return ids


def _chat_anthropic(user, system=None):
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    _stats_call("anthropic", len(system or EXTRACTION_SYSTEM) + len(user))
    msg = client.messages.create(model=model_for("anthropic"), max_tokens=4000,
                                 system=system or EXTRACTION_SYSTEM,
                                 messages=[{"role": "user", "content": user}])
    text = msg.content[0].text
    u = getattr(msg, "usage", None)
    _stats_usage({"prompt_tokens": getattr(u, "input_tokens", 0),
                  "completion_tokens": getattr(u, "output_tokens", 0)} if u else None, len(text))
    return text


# Default model chain, best→most-reliable. Every call walks this until one link
# succeeds, so a slow/rate-limited/JSON-flaky provider never fails the run: the strong,
# no-rate-limit models lead; NIM's clean-JSON llama is the backstop that's almost never
# empty. Links whose API key is absent are dropped. Override with BRIEF_MODEL_CHAIN.
_DEFAULT_CHAIN = [
    ("cerebras", "gpt-oss-120b"),                  # fast, no rate limit, strong generation
    ("cerebras", "zai-glm-4.7"),                   # 2nd Cerebras model (different failure mode)
    ("groq",     "openai/gpt-oss-120b"),           # same strong model, different host
    ("groq",     "llama-3.3-70b-versatile"),       # fast, clean JSON
    ("nim",      "meta/llama-3.3-70b-instruct"),   # reliable clean-JSON backstop
    ("nim",      "nvidia/llama-3.1-nemotron-70b-instruct"),
]
_KEY_ENV = {"cerebras": "CEREBRAS_API_KEY", "groq": "GROQ_API_KEY", "nim": "NVIDIA_API_KEY",
            "gemini": "GEMINI_API_KEY", "openai": "OPENAI_API_KEY", "anthropic": "ANTHROPIC_API_KEY",
            "ollama": None}
_CHAIN_LOGGED = False
# Circuit-breaker: a link that 429s is skipped for this many seconds so subsequent calls
# settle onto a free link instead of re-failing the saturated lead on every single call.
_LINK_COOLDOWN = {}          # "provider:model" -> monotonic deadline
_COOLDOWN_SECS = float(os.environ.get("BRIEF_LINK_COOLDOWN", "45"))


def _cooldown(provider, model):
    import time
    _LINK_COOLDOWN[f"{provider}:{model}"] = time.monotonic() + _COOLDOWN_SECS


def _provider_for_model(m: str) -> "str | None":
    """Best-guess host for an explicitly-pinned model id (a wrong guess self-heals — the
    chain falls through to the next link if the pinned link errors)."""
    if m.startswith(("nvidia/", "meta/")):
        return "nim"
    if m in ("gpt-oss-120b", "zai-glm-4.7"):
        return "cerebras"
    if m.startswith(("openai/", "meta-llama/", "qwen/")) or m.endswith("-versatile"):
        return "groq"
    return None


def _model_chain(model=None) -> list:
    """Ordered [(provider, model)] to try, best first. BRIEF_MODEL_CHAIN overrides the
    default; an explicit BRIEF_PROVIDER (and model=) is honoured as the FIRST link, with
    the rest kept as fallback so a pinned provider still never fails the run."""
    global _CHAIN_LOGGED
    env_chain = os.environ.get("BRIEF_MODEL_CHAIN", "").strip()
    if env_chain:
        chain = [tuple(t.strip().split(":", 1)) for t in env_chain.split(",")
                 if ":" in t]
    else:
        chain = [(p, m) for (p, m) in _DEFAULT_CHAIN
                 if p == "ollama" or os.environ.get(_KEY_ENV.get(p) or "")]
    # Honour an explicit pin (BRIEF_PROVIDER / model=) as the lead link, keep fallback.
    pin_p = os.environ.get("BRIEF_PROVIDER", "").lower().strip()
    pin_m = model or os.environ.get("BRIEF_MODEL", "")
    if pin_p:
        lead = (pin_p, pin_m or model_for(pin_p))
        chain = [lead] + [l for l in chain if l != lead]
    elif model:
        prov = _provider_for_model(model) or (chain[0][0] if chain else resolve_provider())
        lead = (prov, model)
        chain = [lead] + [l for l in chain if l != lead]
    if not _CHAIN_LOGGED and chain:
        print("[i] model chain: " + " → ".join(f"{p}:{m}" for p, m in chain), file=sys.stderr)
        _CHAIN_LOGGED = True
    # Drop links that are mid-cooldown (recently 429'd); if every link is cooling, keep the
    # full chain rather than deadlock — something is better than heuristic.
    import time
    now = time.monotonic()
    live = [l for l in chain if _LINK_COOLDOWN.get(f"{l[0]}:{l[1]}", 0.0) <= now]
    return live or chain


def _call_link(provider: str, model: str, user, system=None, max_tokens=None,
               json_mode=False) -> "str | None":
    """Call exactly ONE (provider, model) link. Raises on failure so the chain advances."""
    if provider == "anthropic":
        return _chat_anthropic(user, system=system)
    cfg = PROVIDERS.get(provider)
    if not cfg:
        raise RuntimeError(f"unknown provider '{provider}'")
    default_base, key_env, _ = cfg
    key = os.environ.get(key_env) if key_env else "ollama"
    if key_env and not key:
        raise RuntimeError(f"{key_env} not set")
    return _chat_openai_compatible(default_base, key, model, user,
                                   provider_label=f"{provider}:{model}", system=system,
                                   max_tokens=max_tokens, json_mode=json_mode)


def _stats_logical():
    """One LOGICAL call (one _chat/_json_call) may cost several link attempts when the
    chain fails over — 'calls' counts attempts, this counts intent."""
    if not _LLM_STATS:
        _stats_reset()
    with _STATS_LOCK:
        _LLM_STATS["logical_calls"] = _LLM_STATS.get("logical_calls", 0) + 1


def _chat(user, system=None, model=None, max_tokens=None):
    """Provider-agnostic chat. Walks the model chain (see _model_chain) and returns the
    first link's raw text, or None if every link failed (callers fall back to heuristic)."""
    _stats_logical()
    for provider, m in _model_chain(model):
        try:
            raw = _call_link(provider, m, user, system=system, max_tokens=max_tokens)
            if raw:
                return raw
        except _RateLimited:                     # saturated — cool it down so later calls skip it
            _cooldown(provider, m)
            print(f"[i] link {provider}:{m} rate-limited; cooling {int(_COOLDOWN_SECS)}s, next link…",
                  file=sys.stderr)
        except Exception as e:                   # network/auth/HTTP — try the next link
            print(f"[i] link {provider}:{m} failed ({e.__class__.__name__}); next link…",
                  file=sys.stderr)
    return None


def _json_call(user, system=None, retries=1, model=None, accept=None, max_tokens=None):
    """Chat call that must return JSON, with provider juggling: each chain link gets up to
    `retries`+1 tries; unparseable output (or one rejected by `accept`) advances to the next
    link. Returns the first usable object, or None if the whole chain is exhausted."""
    _stats_logical()
    for provider, m in _model_chain(model):
        for attempt in range(retries + 1):
            try:
                raw = _call_link(provider, m, user, system=system, max_tokens=max_tokens,
                                 json_mode=True)
            except _RateLimited:
                _cooldown(provider, m)
                print(f"[i] link {provider}:{m} rate-limited; cooling {int(_COOLDOWN_SECS)}s, next link…",
                      file=sys.stderr)
                break
            except Exception as e:
                print(f"[i] link {provider}:{m} failed ({e.__class__.__name__}); next link…",
                      file=sys.stderr)
                break                            # this link is down — go to the next one
            if not raw:
                break
            raw = re.sub(r"^```(?:json)?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
            obj = _loads_lenient(raw)
            if obj is not None and (accept is None or accept(obj)):
                return obj
            if attempt < retries:
                print(f"[i] {provider}:{m} unclean/unusable JSON; retrying once.", file=sys.stderr)
        # link exhausted → fall through to the next provider in the chain
    print("[i] no provider in the chain returned usable JSON.", file=sys.stderr)
    return None


def extract_llm(text, schema):
    """Provider-agnostic extractor. Walks the model chain (BRIEF_MODEL_CHAIN); a link that
    returns valid-but-empty JSON is skipped so extraction lands on a clean-JSON model.
    Returns a dict, or None to fall back to heuristic."""
    obj = _json_call(_user_msg(text, schema), max_tokens=MAXTOK_EXTRACT,
                     accept=lambda o: isinstance(o, dict) and isinstance(o.get("fields"), dict)
                     and bool(o["fields"]))
    out = _normalize_llm(obj) if obj is not None else None
    # Shape guard: the lenient parser can fish an inner {...} out of truncated
    # output — an "extraction" with no fields must not count as a success.
    if not (isinstance(out, dict) and isinstance(out.get("fields"), dict)
            and out["fields"]):
        if obj is not None:
            print("[i] LLM JSON had no usable 'fields'.", file=sys.stderr)
        print("[i] Falling back to heuristic extraction.", file=sys.stderr)
        return None
    return out


def _loads_lenient(raw):
    """Reasoning models sometimes wrap the JSON in prose. Try a clean parse,
    then every balanced {...} span (largest first), each with a trailing-comma
    repair pass. Returns dict/list or None."""
    def _try(s):
        for candidate in (s, re.sub(r",\s*([}\]])", r"\1", s)):
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                continue
        return None

    obj = _try(raw)
    if obj is not None:
        return obj
    spans = []
    start = raw.find("{")
    while start != -1:
        depth, in_str, esc = 0, False, False
        for i in range(start, len(raw)):
            ch = raw[i]
            if in_str:
                if esc:        esc = False
                elif ch == "\\": esc = True
                elif ch == '"': in_str = False
            elif ch == '"':    in_str = True
            elif ch == "{":    depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    spans.append(raw[start:i + 1])
                    break
        start = raw.find("{", start + 1)
    for span in sorted(set(spans), key=len, reverse=True):
        obj = _try(span)
        if isinstance(obj, dict):
            return obj
    return None


# LLM field name (free-form) -> the canonical key the rest of the pipeline keys on.
# The model invents synonyms; without this, Loop 2 shaping and the renderer (which
# only walk canonical keys) silently drop genuinely-extracted content.
FIELD_ALIASES = {
    "audience": "target_audience", "target": "target_audience",
    "mandatory_requirements": "mandatories", "mandatory": "mandatories",
    "must_haves": "mandatories", "requirements": "mandatories",
    "competitors": "competitors_market", "competition": "competitors_market",
    "competitive_landscape": "competitors_market",
    "problem": "business_problem", "challenge": "business_problem",
    "business_challenge": "business_problem",
    "background": "background_context", "context": "background_context",
    "goal": "objective", "goals": "objective", "objectives": "objective",
    "kpis": "success_metrics", "metrics": "success_metrics", "success": "success_metrics",
    "timing": "timeline", "deadline": "timeline", "key_dates": "timeline",
    "deliverable": "deliverables",
    "tone": "tone_and_brand", "brand": "tone_and_brand",
    "brand_guidelines": "tone_and_brand", "tone_of_voice": "tone_and_brand",
    "message": "key_message", "key_messages": "key_message",
    "proposition": "key_message", "single_minded_proposition": "key_message",
    "proof": "proof_points", "proofs": "proof_points", "rtbs": "proof_points",
    "reasons_to_believe": "proof_points", "support": "proof_points",
    "evaluation": "evaluation_criteria", "judging_criteria": "evaluation_criteria",
    "assessment_criteria": "evaluation_criteria",
    "strategy": "strategic_angle", "strategic_direction": "strategic_angle",
    "anti_audience": "anti_target", "non_target": "anti_target",
}


def _canonicalize_fields(fields):
    """Rename known LLM synonyms to canonical keys; merge if both exist."""
    if not isinstance(fields, dict):
        return fields
    out = {}
    for k, v in fields.items():
        ck = FIELD_ALIASES.get(k, k)
        if ck in out:                                # merge collisions into a list
            cur = out[ck]
            out[ck] = (cur if isinstance(cur, list) else [cur]) + \
                      (v if isinstance(v, list) else [v])
        else:
            out[ck] = v
    return out


def _coerce_open_qs(qs):
    """LLMs sometimes return open_questions as bare strings — wrap into the
    {question, ...} dict shape the ledger and renderer consume."""
    out = []
    for q in (qs if isinstance(qs, list) else []):
        if isinstance(q, str) and q.strip():
            out.append({"question": q.strip()})
        elif isinstance(q, dict):
            out.append(q)
    return out


def _normalize_llm(d):
    """Tolerate two output shapes. The prompt asks for the flat
    {fields, how_to_win, open_questions}, but models fed the full schema often
    mirror it instead (fields under loop1_capture, open_questions under
    loop2_brief). Unwrap either into the flat shape the pipeline consumes, and
    canonicalize field keys either way."""
    if not isinstance(d, dict):
        return d
    if "fields" in d:                                # already flat
        d["fields"] = _canonicalize_fields(d.get("fields"))
        d["open_questions"] = _coerce_open_qs(d.get("open_questions"))
        return d
    l1 = d.get("loop1_capture", {}) if isinstance(d.get("loop1_capture"), dict) else {}
    l2 = d.get("loop2_brief", {}) if isinstance(d.get("loop2_brief"), dict) else {}
    fields = l1.get("fields") or d.get("capture") or {}
    how_to_win = l1.get("how_to_win") or d.get("how_to_win") or {}
    open_qs = (l2.get("open_questions") or d.get("open_questions")
               or l1.get("open_questions") or [])
    return {"fields": _canonicalize_fields(fields),
            "how_to_win": how_to_win, "open_questions": _coerce_open_qs(open_qs)}


# ---------------------------------------------------------------------------
# 4. NO-LOSS LEDGER
# ---------------------------------------------------------------------------

def _norm(s):
    if isinstance(s, (list, tuple)):  # LLMs sometimes return a list of values
        s = " ".join(str(x) for x in s)
    return re.sub(r"[^a-z0-9 ]", "", str(s).lower())


def build_ledger(segments, used_idx, fields, source_name, how_to_win=None, open_qs=None):
    if used_idx is None:
        used_idx = set()
        quotes = []
        # 1) verbatim source_quotes (and values) from the captured fields
        for v in fields.values():
            for it in (v if isinstance(v, list) else [v]):
                if isinstance(it, dict):
                    for k in ("source_quote", "value"):
                        if it.get(k):
                            quotes.append(_norm(it[k]))
        # 2) evidence quotes used in how_to_win count as "captured" too
        for items in (how_to_win or {}).values():
            for it in (items or []):
                if isinstance(it, dict):
                    for k in ("evidence", "point"):
                        if it.get(k):
                            quotes.append(_norm(it[k]))
        # 3) text the open questions were derived from
        for q in (open_qs or []):
            if q.get("question"):
                quotes.append(_norm(q["question"]))
        quotes = [q for q in quotes if q]
        # A segment counts as mapped on a substring hit, or when most of its
        # content words appear in one quote — verbatim-only matching capped
        # coverage ~50% even on good extractions (quotes get lightly rephrased).
        quote_words = [(q, set(q.split())) for q in quotes]
        for i, seg in enumerate(segments):
            ns = _norm(seg)
            sw = set(ns.split())
            for q, qw in quote_words:
                if ns and (ns in q or q in ns):
                    used_idx.add(i); break
                if len(sw) >= 4 and len(sw & qw) / len(sw) >= 0.7:
                    used_idx.add(i); break
    unmapped = [{"segment": s, "source_ref": source_name}
                for i, s in enumerate(segments) if i not in used_idx]
    total = len(segments); mapped = total - len(unmapped)
    return {"total_segments": total, "mapped_segments": mapped,
            "coverage_pct": round(100 * mapped / total, 1) if total else 0.0,
            "unmapped": unmapped}


# ---------------------------------------------------------------------------
# 5. SELF-CRITIQUE (the REVIEW step of each loop)
# ---------------------------------------------------------------------------

# core field -> (fallback fields that also satisfy it, why it matters, priority)
# why_it_matters carries the BetterBriefs evidence — these questions go back to
# the client, and the stats are the ammunition for asking them.
CORE_FIELDS = {
    "business_problem": ([], "We can't position the work without the real problem.", "blocker"),
    "objective": (["success_metrics"], "Objectives are the most critical yet most "
                  "poorly defined element of a brief — 61% of marketers and 71% of "
                  "agencies rank them #1 (BetterBriefs).", "blocker"),
    "target_audience": ([], "65% of agencies can't picture the target from the briefs "
                        "they get; if we can't picture them, neither can creatives "
                        "(BetterBriefs).", "blocker"),
    "key_message": ([], "A good brief lands ONE single-minded message backed by proof "
                    "points — not a shopping list (BetterBriefs).", "important"),
    "evaluation_criteria": (["success_metrics"], "Only 30% of clients define how work "
                            "will be judged, and 88% of agencies are unclear on it — "
                            "agreeing criteria upfront prevents subjective rounds of "
                            "rework (BetterBriefs).", "important"),
    "budget": ([], "Budget, objectives and audience must interlock — scope and "
               "ambition depend on the money (BetterBriefs).", "important"),
    "timeline": ([], "Drives feasibility and the pitch date.", "important"),
    "success_metrics": (["objective"], "Objectives need benchmarks and a time stamp; "
                        "Loop 7 (IPA QA) can't score without KPIs.", "important"),
    "mandatories": ([], "Missing mandatories = legal/brand risk downstream.", "important"),
    "decision_makers": ([], "We win the room by knowing who decides — 62% of marketers "
                        "vs 43% of agencies say the right people sign off "
                        "(BetterBriefs).", "important"),
}

# When the client gave no evaluation criteria, ask with Orlando Wood's three
# tests rather than a bare "what are the criteria?".
EVAL_CRITERIA_QUESTION = (
    "How will the work be evaluated? Can we agree criteria that (1) connect to "
    "real-world business outcomes, (2) give oxygen to the creative idea rather "
    "than reduce its impact, and (3) indicate whether the work builds mental "
    "availability or fame?")


def review_loop1(ledger, fields):
    flags = []
    if ledger["coverage_pct"] < 85:
        flags.append(f"Coverage {ledger['coverage_pct']}% < 85% — "
                     f"{len(ledger['unmapped'])} segments need a home (re-pass).")
    assumptions = sum(1 for v in fields.values()
                      for it in (v if isinstance(v, list) else [v])
                      if isinstance(it, dict) and it.get("status") == "assumption")
    if assumptions:
        flags.append(f"{assumptions} value(s) are assumptions — verify with client.")
    return {"passed": not flags, "flags": flags}


def review_loop2(loop2):
    missing = [k for k in ("problem", "objective", "audience")
               if not loop2.get(k) or not loop2[k].get("value")]
    flags = [f"Agency brief missing: {m}" for m in missing]
    return {"passed": not flags, "flags": flags}


# ---------------------------------------------------------------------------
# 5b. BETTERBRIEFS SCORECARD — quality of the client brief, not just presence
# ---------------------------------------------------------------------------

SCORECARD_DIMENSIONS = ("objectives_quality", "audience_vividness",
                        "single_minded_message", "evaluation_criteria",
                        "budget_interlock", "strategic_clarity", "language")

# Demographic clichés = "as sure a sign as any that you have not got a strategy".
AUDIENCE_CLICHES = re.compile(
    r"\b(millennials?|gen\s*[zxy]|boomers?|everyone|general (?:public|population)"
    r"|adults?\s*\d{2}\s*[-–]\s*\d{2}|all (?:adults|consumers))\b", re.I)


def _dim(dimension, verdict, evidence, fix=""):
    return {"dimension": dimension, "verdict": verdict,
            "evidence": evidence, "fix": fix}


def scorecard_heuristic(fields):
    """No-API scorecard: presence + cheap quality cues only. The LLM judge is
    the real test; this keeps the pipeline keyless-safe."""
    def items(key):
        v = fields.get(key)
        return [x for x in (v if isinstance(v, list) else [v])
                if isinstance(x, dict) and x.get("value")]

    dims = []
    objs = items("objective")
    if not objs and not items("success_metrics"):
        dims.append(_dim("objectives_quality", "missing", "no objective captured",
                         "Ask for a handful of benchmarked, time-stamped objectives."))
    elif len(objs) > 5:
        dims.append(_dim("objectives_quality", "vague",
                         f"{len(objs)} objectives captured — more than a handful",
                         "Adding objectives dramatically reduces the odds any works."))
    else:
        txt = " ".join(str(o.get("value")) for o in objs + items("success_metrics"))
        timed = bool(re.search(r"\b(20\d\d|q[1-4]|by \w+|\d+\s*(%|pts?|points))", txt, re.I))
        dims.append(_dim("objectives_quality", "pass" if timed else "vague",
                         "benchmark/time-stamp found" if timed
                         else "no benchmark or time stamp detected",
                         "" if timed else "Objectives need benchmarks and a time stamp."))

    aud = _val(fields, "target_audience")
    if not aud:
        dims.append(_dim("audience_vividness", "missing", "no audience captured",
                         "Ask for a vivid picture: demographics + psychographics + needs."))
    elif AUDIENCE_CLICHES.search(aud) or len(aud) < 40:
        dims.append(_dim("audience_vividness", "vague", aud[:120],
                         "Demographic clichés signal no strategy — push for a portrait."))
    else:
        dims.append(_dim("audience_vividness", "pass", aud[:120]))

    dims.append(_dim("single_minded_message",
                     "pass" if _val(fields, "key_message") else "missing",
                     _val(fields, "key_message") or "no key message captured",
                     "" if _val(fields, "key_message")
                     else "A good brief lands ONE message with proof points."))
    dims.append(_dim("evaluation_criteria",
                     "pass" if items("evaluation_criteria") else "missing",
                     "; ".join(str(x.get("value")) for x in items("evaluation_criteria"))[:120]
                     or "no criteria captured", ""))
    dims.append(_dim("budget_interlock",
                     "pass" if _val(fields, "budget") else "missing",
                     _val(fields, "budget") or "no budget captured",
                     "" if _val(fields, "budget")
                     else "Budget, objectives and audience must interlock."))
    dims.append(_dim("strategic_clarity",
                     "pass" if _val(fields, "strategic_angle") else "vague",
                     _val(fields, "strategic_angle") or "no strategic angle captured",
                     "" if _val(fields, "strategic_angle")
                     else "Don't leave strategy for the creative process to discover."))
    dims.append(_dim("language", "pass", "not assessed in heuristic mode"))
    return {"dimensions": dims,
            "single_mindedness": {"verdict": "single", "split_into": []},
            "summary": "Heuristic scorecard — run with an LLM key for the real judge.",
            "mode": "heuristic"}


def score_betterbriefs(text, fields):
    """LLM judge against the BetterBriefs rubric; heuristic fallback."""
    obj = _json_call(f"CLIENT BRIEF:\n\"\"\"\n{_clip_brief(text)}\n\"\"\"", system=SCORECARD_SYSTEM,
                     max_tokens=MAXTOK_EXTRACT)
    if not isinstance(obj, dict) or not isinstance(obj.get("dimensions"), list):
        return scorecard_heuristic(fields)
    dims = []
    for d in obj["dimensions"]:
        if isinstance(d, dict) and d.get("dimension") in SCORECARD_DIMENSIONS:
            dims.append(_dim(d["dimension"],
                             d.get("verdict") if d.get("verdict") in
                             ("pass", "vague", "missing") else "vague",
                             str(d.get("evidence") or ""), str(d.get("fix") or "")))
    scored = {x["dimension"] for x in dims}
    for missing in SCORECARD_DIMENSIONS:          # judge skipped one — make it visible
        if missing not in scored:
            dims.append(_dim(missing, "vague", "not scored by judge"))
    sm = obj.get("single_mindedness") or {}
    if not isinstance(sm, dict):
        sm = {}
    split = sm.get("split_into")
    return {"dimensions": dims,
            "single_mindedness": {
                "verdict": "multiple" if sm.get("verdict") == "multiple" else "single",
                "split_into": [str(s) for s in split] if isinstance(split, list) else []},
            "summary": str(obj.get("summary") or ""), "mode": "llm"}


# ---------------------------------------------------------------------------
# 6. LOOP 2 — shape the capture into a first-round agency brief
# ---------------------------------------------------------------------------

def _val(fields, key):
    v = fields.get(key)
    if isinstance(v, list):
        parts = []
        for x in v:
            if isinstance(x, dict) and x.get("value"):
                parts.append(str(x["value"]))
            elif isinstance(x, str) and x.strip():
                parts.append(x.strip())
        return "; ".join(parts) or None
    return v.get("value") if isinstance(v, dict) else None


def shape_loop2(fields, llm_open_qs):
    """Map Loop-1 capture into an agency-brief shape + open questions.
    Deterministic so it runs without an API; LLM open-questions used if present."""
    def slot(text):
        return {"value": text, "status": "fact" if text else "gap"}

    scope_bits = [b for b in (_val(fields, "deliverables"), _val(fields, "budget"),
                              _val(fields, "timeline")) if b]
    loop2 = {
        "problem": slot(_val(fields, "business_problem")),
        "objective": slot(_val(fields, "objective") or _val(fields, "success_metrics")),
        "audience": slot(_val(fields, "target_audience")),
        # BetterBriefs slots: the single-minded message (+ proof), how the work
        # will be judged, and the strategic sacrifice (who/what we're NOT for).
        "key_message": slot(_val(fields, "key_message")),
        "evaluation_criteria": slot(_val(fields, "evaluation_criteria")),
        "not_doing": slot(_val(fields, "anti_target")),
        "scope": slot(" · ".join(scope_bits) if scope_bits else None),
    }

    # open questions = genuine gaps in the core fields (a field counts as present
    # if it OR one of its fallbacks was captured), plus any the LLM surfaced.
    open_qs = list(llm_open_qs or [])
    for key, (fallbacks, why, priority) in CORE_FIELDS.items():
        if _val(fields, key) or any(_val(fields, fb) for fb in fallbacks):
            continue
        question = (EVAL_CRITERIA_QUESTION if key == "evaluation_criteria"
                    else f"What is the {key.replace('_', ' ')}?")
        open_qs.append({
            "question": question, "why_it_matters": why, "priority": priority,
        })
    rank = {"blocker": 0, "important": 1, "nice_to_have": 2}

    def _q_rank(q):
        pr = q.get("priority") if isinstance(q, dict) else None
        return rank.get(pr, 3) if isinstance(pr, str) else 3

    open_qs.sort(key=_q_rank)
    loop2["open_questions"] = open_qs
    return loop2


# ---------------------------------------------------------------------------
# 7. RENDER (Brain markdown mirror)
# ---------------------------------------------------------------------------

FIELD_TITLES = {
    "background_context": "Background / context", "business_problem": "Business problem",
    "objective": "Objective", "target_audience": "Target audience",
    "key_message": "Key message (single-minded)", "proof_points": "Proof points",
    "evaluation_criteria": "Evaluation criteria", "strategic_angle": "Strategic angle",
    "anti_target": "Not for / not targeting",
    "deliverables": "Deliverables", "mandatories": "Mandatories (non-negotiable)",
    "budget": "Budget", "timeline": "Timeline & key dates",
    "success_metrics": "Success metrics / KPIs", "competitors_market": "Competitors & market",
    "tone_and_brand": "Tone & brand", "decision_makers": "Decision-makers",
    "constraints": "Constraints",
}
VERDICT_ICON = {"pass": "✅", "vague": "⚠️", "missing": "❌"}
STATUS_TAG = {"fact": "", "assumption": " _(assumption)_", "gap": " _(gap)_"}


def _fmt(c):
    if not c or c.get("value") in (None, ""):
        return "_not stated_"
    return f"{c['value']}{STATUS_TAG.get(c.get('status', 'fact'), '')}"


# ---------------------------------------------------------------------------
# 6b. LOOPS 3–7 — RAG-grounded strategy
#     The ONLY place retrieval happens. Built from the Loop-2 brief, never from
#     Loop-1 capture, and the rag retriever is imported lazily so the Loop-1
#     path never even touches `rag`. Behind a flag; degrades to a stub if the
#     index hasn't been built.
# ---------------------------------------------------------------------------

# Each loop maps a stage of the strategic process to a query built from the
# brief gist. Retrieval grounds it in the planner playbooks + IPA evidence.
LOOP37_SPECS = [
    ("loop3_research", "Loop 3 · Research & category intelligence",
     lambda g: f"how to research the category, competitors and audience for {g['audience']}; {g['problem']}"),
    ("loop4_insight", "Loop 4 · Human insight & cultural tension",
     lambda g: f"find the human insight and cultural tension for {g['audience']} given {g['problem']}"),
    ("loop5_proposition", "Loop 5 · Single-minded proposition",
     lambda g: f"single-minded proposition and key message to achieve {g['objective']}; {g['key_message']}"),
    ("loop6_substantiation", "Loop 6 · Substantiation & effectiveness evidence",
     lambda g: f"effectiveness evidence and proof a strategy delivers {g['objective']}; how brands grow"),
    ("loop7_qa", "Loop 7 · Strategic QA & decision rules",
     lambda g: f"common mistakes and decision rules to pressure-test {g['key_message']} for {g['objective']}"),
]


def _rerank_hits(query: str, hits: list, k: int) -> list:
    """Second-stage rerank: dense retrieval gives recall (vector-similar), this gives
    precision (actually-relevant). The hosted NVIDIA nv-rerankqa NIM isn't provisioned on
    this key, so we rerank in-house via the model chain — keeps retrieval self-hosted and
    provenance-preserving. Best-effort: falls back to the cosine order on any failure.
    Disable with BRIEF_RERANK=0."""
    if os.environ.get("BRIEF_RERANK", "1") == "0" or len(hits) <= k:
        return hits[:k]
    def _snippet(t):
        return re.sub(r'\\s+', ' ', t)[:200]
    listing = "\n".join(f"[{i}] {h.get('citation','')}: {_snippet(h.get('text',''))}"
                        for i, h in enumerate(hits))
    obj = _json_call(
        f"QUERY: {query}\n\nPASSAGES:\n{listing}",
        system=("You are a retrieval reranker for a strategy brief. Rank the passages by how directly "
                "each one helps answer the QUERY — most useful first. Demote passages that merely share "
                "words but miss the intent. Return ONLY raw JSON: {\"ranking\": [passage indexes, best first]}"),
        retries=1, max_tokens=MAXTOK_JUDGE)
    if isinstance(obj, dict) and isinstance(obj.get("ranking"), list):
        order = [i for i in obj["ranking"] if isinstance(i, int) and 0 <= i < len(hits)]
        order += [i for i in range(len(hits)) if i not in order]   # keep any the judge dropped
        return [hits[i] for i in order][:k]
    return hits[:k]


def _load_retriever():
    """Import the Loops 3–7 retriever lazily. Adds briefing/rag/ to sys.path and
    imports the top-level `retrieve` module — whose own `import rag` then resolves
    to rag/rag.py (not the rag/ dir as a package). Keeps the Loop-1 path rag-free."""
    rag_dir = HERE / "rag"
    if str(rag_dir) not in sys.path:
        sys.path.insert(0, str(rag_dir))
    import retrieve                                # noqa: E402  (rag/retrieve.py)
    return retrieve


def _capsule_text(v) -> str:
    """Plain text from a Loop-2 capsule / Loop-1 field (dict | list | str)."""
    if isinstance(v, dict):
        return str(v.get("value") or "")
    if isinstance(v, list):
        return " · ".join(t for t in (_capsule_text(x) for x in v) if t)
    return str(v or "")


def _brief_gist(loop2, fields) -> dict:
    g = {k: _capsule_text(loop2.get(k))
         for k in ("problem", "objective", "audience", "key_message")}
    g["problem"] = g["problem"] or _val(fields, "business_problem") or ""
    g["objective"] = g["objective"] or _val(fields, "objective") or _val(fields, "success_metrics") or ""
    g["audience"] = g["audience"] or _val(fields, "target_audience") or ""
    g["key_message"] = g["key_message"] or _val(fields, "key_message") or ""
    return g


def _classify_intent(gist, fields) -> str:
    """Lightweight brief-intent label. The full Loop-0 classifier is backlog #2;
    here it just biases the read and is surfaced in the output."""
    blob = (" ".join(gist.values()) + " "
            + " ".join(_capsule_text(fields.get(k)) for k in fields)).lower()
    table = [("tender", ("tender", "rfp", "itt", "procurement", "pitch document")),
             ("media", ("media plan", "media buying", "channel mix", "reach and frequency", "grp")),
             ("btl-event", ("activation", "experiential", " btl", "sampling", "event")),
             ("retail", ("shopper", "in-store", "point of sale", "retail media", "trade")),
             ("creative-campaign", ("campaign", "creative", "advert", "tvc", "film", "launch"))]
    for label, kws in table:
        if any(kw in blob for kw in kws):
            return label
    return "general-strategy"


def _synthesize_loops37(gist, intent, loops) -> str:
    """Ground a short paragraph per loop in the retrieved evidence, citing
    `source › section`. Reuses the existing LLM provider; falls back to an
    evidence-only summary when no provider / no clean JSON. Mutates loops."""
    blocks = []
    for key, d in loops.items():
        ev = "\n".join(f"  - ({e['citation']}) {e['snippet']}" for e in d["evidence"]) \
             or "  (no evidence retrieved)"
        blocks.append(f"### {key} — {d['title']}\n{ev}")
    user = (
        "You are an advertising planning director. Using ONLY the retrieved evidence "
        "below, write one grounded, specific paragraph per loop that applies the "
        "frameworks to THIS brief. Cite the playbooks you use inline as "
        "(source › section), copied exactly. Never invent frameworks or statistics.\n\n"
        f"BRIEF GIST: problem={gist['problem']!r}; objective={gist['objective']!r}; "
        f"audience={gist['audience']!r}; key_message={gist['key_message']!r}; intent={intent}.\n\n"
        f"RETRIEVED EVIDENCE:\n" + "\n\n".join(blocks) + "\n\n"
        "Return JSON only: an object mapping each loop key "
        f"({', '.join(loops)}) to its paragraph string."
    )
    synth_model = os.environ.get("BRIEF_SYNTH_MODEL") or None
    obj = _json_call(user, system="You are a precise strategy planner. Output JSON only.",
                     model=synth_model, max_tokens=MAXTOK_EXTRACT)
    if isinstance(obj, dict):
        wrote = False
        for key, d in loops.items():
            para = obj.get(key)
            if isinstance(para, str) and para.strip():
                d["synthesis"] = para.strip(); wrote = True
        if wrote:
            prov = resolve_provider()
            used = synth_model or model_for(prov)
            return f"llm:{used}" if prov else "llm"
    for d in loops.values():                          # evidence-only fallback
        if d["evidence"]:
            tops = "; ".join(f"{e['framework']} ({e['citation']})" for e in d["evidence"][:3])
            d["synthesis"] = f"Apply, in order of fit: {tops}."
        else:
            d["synthesis"] = "No playbook evidence retrieved for this loop."
    return "evidence-only"


def _loops37_from_digests(loop2, fields) -> dict | None:
    """Digest mode: no vector store, but pack digests (packs_dist/<id>/digest.md,
    written offline by scripts/distil_pack.py) exist. Ground Loops 3–7 on those —
    static per pack rather than query-matched, but the synthesis, citations and
    strategy fill all run. This is the app's default grounding path."""
    digest_dir = HERE / "packs_dist"
    digests = sorted(digest_dir.glob("*/digest.md")) if digest_dir.is_dir() else []
    entries = []
    for d in digests:
        text = d.read_text(errors="replace").strip()
        if text:
            entries.append({
                "citation": f"{d.parent.name} digest",
                "framework": f"{d.parent.name} digest",
                "category": None,
                "score": 0.0,
                "snippet": re.sub(r"\s+", " ", text)[:600],
            })
    if not entries:
        return None
    gist = _brief_gist(loop2, fields)
    intent = _classify_intent(gist, fields)
    loops = {key: {"title": title, "query": "(digest mode — no retrieval)",
                   "evidence": list(entries)}
             for key, title, _q in LOOP37_SPECS}
    synthesis_mode = _synthesize_loops37(gist, intent, loops)
    return {
        "enabled": True,
        "index": "digests:packs_dist",
        "intent": intent,
        "k": 0,
        "gist": gist,
        "loops": loops,
        "sources_used": sorted({e["citation"] for e in entries}),
        "synthesis_mode": synthesis_mode,
    }


def loops_3_7(loop2, fields, k=5, index_dir=None) -> dict:
    """Loops 3–7: classify intent → build queries from the Loop-2 brief → retrieve
    top-k playbooks + effectiveness evidence → ground a short strategy with
    citations. Retrieval-only; degrades to a disabled stub if the index is absent."""
    try:
        retriever = _load_retriever()
    except Exception as e:                            # never crash the run over RAG
        return {"enabled": False,
                "reason": f"retriever import failed: {e.__class__.__name__}: {e}"}
    if not retriever.index_available(index_dir):
        digest_loops = _loops37_from_digests(loop2, fields)
        if digest_loops:
            return digest_loops
        return {"enabled": False,
                "reason": "no retrieval store (rag/index absent, no Qdrant) and no pack "
                          "digests (packs_dist/) — Loops 3–7 skipped."}

    gist = _brief_gist(loop2, fields)
    intent = _classify_intent(gist, fields)

    # Case packs, discovered from the corpus dirs (or packs.lock at runtime) —
    # never a hardcoded list, so adding/removing a pack needs no code change
    # and a pack with no corpus simply cannot exist (the old `effie` bug).
    try:
        from packs import discover_packs
        case_packs = [p for p in discover_packs() if p.kind == "case"]
    except Exception:
        case_packs = []

    def _one_loop(spec):
        """Retrieval + rerank + precedent pull for ONE loop — fully independent given
        the gist, so the five loops run concurrently (network-bound: Qdrant + NIM
        embeddings + optional rerank). ~5× wall-clock cut on this stage."""
        key, title, qfn = spec
        q = re.sub(r"\s+", " ", qfn(gist)).strip()
        seen, evidence = set(), []
        # Over-retrieve for recall, then LLM-rerank down to k for precision.
        # 1.5× is enough headroom — 3× ranked 15 passages to keep 5 (dead tokens).
        pool = retriever.retrieve(q, k=max(int(k * 1.5), 8), index_dir=index_dir)
        for h in _rerank_hits(q, pool, k):
            if h["source"] in seen:
                continue
            seen.add(h["source"])
            evidence.append({
                "citation": h["citation"],
                "framework": h.get("framework") or h["source"],
                "category": h.get("category"),
                "score": h["score"],
                "snippet": re.sub(r"\s+", " ", h["text"])[:280],
            })
        # Pull award-winning PRECEDENT cases from every case pack whose `loops`
        # gate includes this loop (default: insight + substantiation). Which packs
        # exist, their tag, and their per-pack k all come from the pack itself.
        eligible = [p for p in case_packs if p.eligible(key)]
        if eligible:
            case_q = (f"award-winning precedent insight {gist['audience']} {gist['problem']}"
                      if key == "loop4_insight"
                      else f"award-winning effectiveness results proof {gist['objective']}"
                      if key == "loop6_substantiation" else q)
            for pack in eligible:
                for h in retriever.retrieve(case_q, k=pack.k, index_dir=index_dir,
                                            where={"source": pack.tag}):
                    if h["source"] not in seen:
                        seen.add(h["source"])
                        evidence.append({
                            "citation": h["citation"],
                            "framework": h.get("framework") or h["source"],
                            "category": h.get("category"),
                            "score": h["score"],
                            "snippet": re.sub(r"\s+", " ", h["text"])[:280],
                        })
        return key, {"title": title, "query": q, "evidence": evidence}

    from concurrent.futures import ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=len(LOOP37_SPECS)) as pool_ex:
        results = dict(pool_ex.map(_one_loop, LOOP37_SPECS))
    # Preserve the canonical loop order regardless of completion order.
    loops = {key: results[key] for key, _t, _q in LOOP37_SPECS}
    citations_all = [e["citation"] for d in loops.values() for e in d["evidence"]]

    synthesis_mode = _synthesize_loops37(gist, intent, loops)
    return {
        "enabled": True,
        "index": (f"qdrant:{os.environ.get('QDRANT_COLLECTION', 'napkin_rag')}"
                  if os.environ.get("RAG_STORE", "").lower().strip() == "qdrant"
                  else str(index_dir or getattr(retriever, "DEFAULT_INDEX", HERE / "rag" / "index"))),
        "intent": intent,
        "k": k,
        "gist": gist,
        "loops": loops,
        "sources_used": sorted(set(citations_all)),
        "synthesis_mode": synthesis_mode,
    }


def render_loops37(L, brief):
    """Append the Loops 3–7 markdown section. No-op when the stage didn't run
    (flag off), so Loop-1/Loop-2 output stays byte-for-byte identical."""
    s = brief.get("loops3_7")
    if not s:
        return
    L.append("## Loops 3–7 · RAG-grounded strategy  ")
    if not s.get("enabled"):
        L.append(f"_skipped — {s.get('reason', '')}_\n")
        return
    L.append(f"_intent: {s['intent']} · retrieval: nv-embedqa-e5-v5 · synthesis: {s['synthesis_mode']}_\n")
    for d in s["loops"].values():
        L.append(f"### {d['title']}\n")
        if d.get("synthesis"):
            L.append(f"{d['synthesis']}\n")
        if d["evidence"]:
            L.append("_Grounded in:_")
            for e in d["evidence"]:
                cat = f" · {e['category']}" if e.get("category") else ""
                L.append(f"- **{e['framework']}** ({e['citation']}{cat}) — {e['snippet']}")
            L.append("")
        else:
            L.append("_No playbook evidence retrieved for this loop._\n")
    if s.get("sources_used"):
        L.append(f"_Sources cited: {len(s['sources_used'])} playbook sections._\n")


def render_client_brief(brief) -> str:
    """The DELIVERABLE — only the final brief. Assembles the Golden Brief (facts +
    generated strategy) into a clean one-pager: no loop labels, no provenance tags,
    no ledgers, no scorecard, no 'Grounded in' citations. All of that machinery lives
    in review.md. This is what a creative director actually reads."""
    m = brief["meta"]
    gf = (brief.get("loop2_golden") or {}).get("fields", {}) or {}
    l2 = brief.get("loop2_brief", {}) or {}
    title = m.get("project") or m.get("client") or "Client brief"

    def gv(fid):                      # golden value, else loop-2 fallback for the FACTS only
        f = gf.get(fid)
        v = f.get("value") if isinstance(f, dict) else f
        if v:
            return v
        # Strategy fields (insight, smp, reasons_to_believe, desired_response) must NEVER
        # fall back to a loop-2 value: if generation didn't clear the rubric the field is a
        # real gap (and carries an open question). Falling back would re-show the masterbrand
        # line while also flagging "agree the SMP" — the contradiction. Facts may fall back.
        fb = l2.get({"background": "problem", "objectives": "objective",
                     "audience": "audience"}.get(fid, ""))
        return fb.get("value") if isinstance(fb, dict) else fb   # loop-2 fields are {value,status}

    L = [f"# {title} — Brief", ""]
    TBD = "_To be agreed — see open questions._"

    def text_section(heading, value):
        L.append(f"## {heading}")
        L.append(str(value) if value else TBD)
        L.append("")

    text_section("Background", gv("background"))

    obj = gv("objectives")
    L.append("## Objectives")
    if isinstance(obj, dict):
        for k, lab in (("commercial", "Commercial"), ("behavioural", "Behavioural"),
                       ("attitudinal", "Attitudinal")):
            if obj.get(k):
                L.append(f"- **{lab}:** {obj[k]}")
    elif obj:
        L.append(str(obj))
    else:
        L.append(TBD)
    L.append("")

    text_section("Audience", gv("audience"))
    text_section("Competitor context", gv("competitor_context"))
    text_section("The insight", gv("insight"))
    text_section("Single-minded proposition", gv("smp"))

    rtb = gv("reasons_to_believe")
    L.append("## Reasons to believe")
    if isinstance(rtb, list) and rtb:
        L += [f"- {r if isinstance(r, str) else (r.get('value') if isinstance(r, dict) else r)}"
              for r in rtb]
    elif rtb:
        L.append(str(rtb))
    else:
        L.append(TBD)
    L.append("")

    dr = gv("desired_response")
    L.append("## Desired response")
    if isinstance(dr, dict):
        for k, lab in (("think", "Think"), ("feel", "Feel"), ("do", "Do")):
            if dr.get(k):
                L.append(f"- **{lab}:** {dr[k]}")
    elif dr:
        L.append(str(dr))
    else:
        L.append(TBD)
    L.append("")

    text_section("Tone & world", gv("tone_world_assets"))
    text_section("Budget & scope", gv("budget_scope"))
    text_section("Mandatories", gv("mandatories"))

    oqs = l2.get("open_questions") or []
    if oqs:
        L.append("## Open questions to resolve before research")
        seen = set()
        for q in oqs:
            txt = q if isinstance(q, str) else (q.get("question") or q.get("value") or "")
            key = re.sub(r"[^a-z0-9]+", " ", txt.lower()).strip()   # dedupe near-identical questions
            if not key or key in seen:
                continue
            seen.add(key)
            pr = "" if isinstance(q, str) else (f"**[{q.get('priority')}]** " if q.get("priority") else "")
            L.append(f"- {pr}{txt}")
        L.append("")
    return "\n".join(L).strip() + "\n"


def render_markdown(brief):
    m, l1, l2 = brief["meta"], brief["loop1_capture"], brief["loop2_brief"]
    led = l1["no_loss_ledger"]
    title = m.get("project") or m.get("client") or "Client brief"
    L = [f"# Brief — {title}",
         f"_briefing tool v{m['parser_version']} · {m['parsed_at']} · "
         f"mode: {m['extraction_mode']}_\n"]

    L.append("## Loop 1 · Faithful capture  \n_IPA: background + objectives · no RAG_\n")
    f = l1["fields"]
    # Canonical fields first (ordered), then any extra LLM fields — never drop content.
    extra = [k for k in f if k not in FIELD_TITLES and k not in ("client", "project")]
    for key, tit in list(FIELD_TITLES.items()) + [(k, k.replace("_", " ").title()) for k in extra]:
        if key not in f:
            continue
        v = f[key]
        if isinstance(v, list):
            if v:
                L.append(f"**{tit}**\n")
                L += [f"- {_fmt(it)}" for it in v]; L.append("")
        else:
            L.append(f"**{tit}** — {_fmt(v)}\n")

    L.append("### Win-rules (what the brief reveals)\n")
    htw = l1["how_to_win"]
    titles = {"stated_evaluation_criteria": "How they'll judge us",
              "unstated_needs": "Unstated needs", "likely_landmines": "Landmines",
              "winning_themes": "Recurring themes", "proof_required": "Proof expected"}
    if any(htw.get(k) for k in titles):
        for k, t in titles.items():
            if htw.get(k):
                L.append(f"**{t}**\n")
                for it in htw[k]:
                    if not isinstance(it, dict):
                        L.append(f"- {it}"); continue
                    # LLM uses value/source_quote; heuristic uses point/evidence.
                    point = it.get("point") or it.get("value") or it.get("text") or ""
                    src = it.get("evidence") or it.get("source_quote")
                    ev = f"  \n  ↳ _{src}_" if src else ""
                    L.append(f"- {point}{ev}")
                L.append("")
    else:
        L.append("_Run with an LLM key for the full win-rules read._\n")

    r1 = l1["review"]
    L.append(f"### Loop 1 self-review — {'✅ pass' if r1['passed'] else '⚠️ needs a pass'}\n")
    L += [f"- {x}" for x in r1["flags"]] or ["- clean"]
    L.append(f"\n**No-loss ledger:** {led['coverage_pct']}% "
             f"({led['mapped_segments']}/{led['total_segments']} mapped)\n")
    if led["unmapped"]:
        L.append("_Review queue (nothing dropped silently):_\n")
        L += [f"- {u['segment']}" for u in led["unmapped"]]; L.append("")

    sc = brief.get("betterbriefs_scorecard")
    if sc:
        L.append("### BetterBriefs scorecard — quality of the client brief  \n"
                 f"_rubric: reference/betterbriefs · judge: {sc.get('mode')}_\n")
        L.append("| Dimension | Verdict | Evidence / fix |")
        L.append("|---|---|---|")
        for d in sc["dimensions"]:
            note = d["evidence"] + (f" → _{d['fix']}_" if d.get("fix") else "")
            L.append(f"| {d['dimension'].replace('_', ' ')} "
                     f"| {VERDICT_ICON.get(d['verdict'], '')} {d['verdict']} "
                     f"| {note.replace('|', '/').replace(chr(10), ' ')} |")
        L.append("")
        sm = sc.get("single_mindedness") or {}
        if sm.get("verdict") == "multiple":
            L.append("**⚠️ Not single-minded — one brief = one strategy. Split into:**\n")
            L += [f"- {s}" for s in sm.get("split_into", [])]; L.append("")
        if sc.get("summary"):
            L.append(f"_{sc['summary']}_\n")

    L.append("## Loop 2 · First-round agency brief  \n_IPA: objective + role_\n")
    for k, t in (("problem", "Problem"), ("objective", "Objective"),
                 ("audience", "Audience"), ("key_message", "Key message"),
                 ("evaluation_criteria", "Evaluation criteria"),
                 ("not_doing", "Not doing"), ("scope", "Scope")):
        if k in l2:                                  # old brief_objects lack new slots
            L.append(f"**{t}** — {_fmt(l2[k])}\n")
    L.append("### Open questions (ask before research)\n")
    if l2["open_questions"]:
        for q in l2["open_questions"]:
            if isinstance(q, str):                # LLM returns bare strings
                L.append(f"- {q}"); continue
            pr = f"**[{q.get('priority','')}]** " if q.get("priority") else ""
            why = q.get("why_it_matters") or q.get("why") or ""
            text = q.get("question") or q.get("value") or ""
            L.append(f"- {pr}{text}" + (f"  \n  _why: {why}_" if why else ""))
    else:
        L.append("_None._")
    r2 = l2["review"]
    L.append(f"\n### Loop 2 self-review — {'✅ pass' if r2['passed'] else '⚠️ gaps'}\n")
    L += [f"- {x}" for x in r2["flags"]] or ["- clean"]
    render_loops37(L, brief)                          # no-op unless Loops 3–7 ran
    render_golden_provenance(L, brief)                # per-field RAG citations (review-only)
    return "\n".join(L)


def render_golden_provenance(L, brief):
    """Surface, in review.md only, which RAG sources grounded each generated strategy
    field — the `evidence_ids` we stamp in fill_derivable_fields. Kept OUT of the
    client deliverable by design (a footnoted 'grounded in <case>' reads as harmful);
    this is where the provenance lives for a planner to audit or defend a route."""
    gf = (brief.get("loop2_golden") or {}).get("fields", {}) or {}
    gen = [(fid, f) for fid, f in gf.items()
           if isinstance(f, dict) and f.get("source") == "inferred" and f.get("method", "").startswith("gen:")]
    miss = [(fid, f) for fid, f in gf.items()
            if isinstance(f, dict) and f.get("source") == "missing" and f.get("reason")]
    if not gen and not miss:
        return
    L.append("\n## Generated strategy — RAG provenance  \n"
             "_Review only; never rendered in the client brief._\n")
    for fid, f in gen:
        label = fid.replace("_", " ")
        conf = f.get("confidence")
        cites = ", ".join(f.get("evidence_ids") or []) or "(no IPA cases cited — playbook-grounded)"
        L.append(f"**{label}** — _{f.get('method')}_, confidence {conf}")
        L.append(f"  \n  ↳ grounded in: {cites}")
        if f.get("rationale"):
            L.append(f"  \n  ↳ rationale: _{f['rationale']}_")
        if f.get("judge_note"):
            L.append(f"  \n  ↳ tournament: _{f['judge_note']}_")
        if f.get("alternatives"):
            L.append(f"  \n  ↳ runner-up: {json.dumps(f['alternatives'])[:200]}")
        L.append("")
    for fid, f in miss:
        L.append(f"**{fid.replace('_', ' ')}** — _missing_: {f.get('reason')}")
        L.append("")


# ---------------------------------------------------------------------------
# RICH OUTPUT (docx / pdf via pandoc)
# ---------------------------------------------------------------------------

# pdflatex/xelatex have no colour-emoji glyphs; map the few we emit to ASCII
# so the PDF renders cleanly. (docx keeps the originals — Word has the fonts.)
_PDF_GLYPHS = {"✅": "[PASS]", "⚠️": "[!]", "⚠": "[!]", "❌": "[FAIL]", "✗": "[FAIL]",
               "🟢": "[+]", "🟡": "[~]", "🔴": "[-]", "↳": ">", "•": "-", "→": "->"}


def _find_xelatex() -> str | None:
    return shutil.which("xelatex") or next(
        (p for p in ("/Library/TeX/texbin/xelatex",) if Path(p).exists()), None)


def write_rich_formats(md_text: str, md_file: Path, formats: list[str]) -> list[str]:
    """Emit docx/pdf alongside the markdown one-pager, via pandoc. Returns the
    formats actually written. Degrades with a clear message, never raises."""
    want = [f.strip().lower() for f in formats if f.strip() and f.strip().lower() != "md"]
    if not want:
        return []
    if not shutil.which("pandoc"):
        print("  ⚠️ pandoc not found — skipping docx/pdf  (brew install pandoc)")
        return []
    done = []
    if "docx" in want:
        out = md_file.with_suffix(".docx")
        r = subprocess.run(["pandoc", str(md_file), "-o", str(out)],
                           capture_output=True, text=True)
        if r.returncode == 0:
            done.append("docx")
        else:
            print(f"  ⚠️ docx failed: {r.stderr.strip()[:200]}")
    if "pdf" in want:
        xelatex = _find_xelatex()
        if not xelatex:
            print("  ⚠️ no xelatex engine — skipping pdf  (install BasicTeX/MacTeX)")
        else:
            clean = md_text
            for k, v in _PDF_GLYPHS.items():
                clean = clean.replace(k, v)
            # Stray backslashes (e.g. Windows paths like \ACME leaking from a brief)
            # are undefined LaTeX control sequences and abort the PDF. They're path
            # noise in a deliverable anyway — neutralise to forward slashes.
            clean = clean.replace("\\", "/")
            tmp = md_file.with_name(".brief_pdf_src.md")
            tmp.write_text(clean, encoding="utf-8")
            env = dict(os.environ)
            env["PATH"] = str(Path(xelatex).parent) + os.pathsep + env.get("PATH", "")
            out = md_file.with_suffix(".pdf")
            r = subprocess.run(
                ["pandoc", str(tmp), "-o", str(out),
                 "--pdf-engine=xelatex", "-V", "geometry:margin=2cm"],
                capture_output=True, text=True, env=env)
            tmp.unlink(missing_ok=True)
            if r.returncode == 0:
                done.append("pdf")
            else:
                print(f"  ⚠️ pdf failed: {r.stderr.strip()[:300]}")
    return done


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def run(path: Path | None, client=None, project=None, loops37=False, golden=False,
        raw_text: str | None = None, source_name: str | None = None) -> dict:
    _stats_reset()
    _t_run0 = dt.datetime.now()
    schema = json.loads((HERE / "brief_object.schema.json").read_text())
    if raw_text is not None:                       # pasted email / stdin / --text
        text = raw_text
        src_name = source_name or "pasted-input"
    else:
        text, _mime = ingest(path)
        src_name = path.name
    segs = segment(text)

    provider = resolve_provider()
    llm = extract_llm(text, schema)
    if llm:
        fields = llm.get("fields", {}); how_to_win = llm.get("how_to_win", {})
        llm_oqs = llm.get("open_questions", []); used = None
        mode = f"{provider}:{model_for(provider)}"
    else:
        fields, used = extract_heuristic(segs); how_to_win = {}; llm_oqs = []
        mode = "heuristic"

    loop2 = shape_loop2(fields, llm_oqs)
    loop2["review"] = review_loop2(loop2)

    ledger = build_ledger(segs, used, fields, src_name, how_to_win, loop2["open_questions"])
    loop1 = {"fields": fields, "how_to_win": how_to_win, "no_loss_ledger": ledger}
    loop1["review"] = review_loop1(ledger, fields)

    scorecard = score_betterbriefs(text, fields)

    out = {
        "meta": {"client": client, "project": project, "source_files": [src_name],
                 "parsed_at": dt.datetime.now().isoformat(timespec="seconds"),
                 "parser_version": PARSER_VERSION, "extraction_mode": mode,
                 "prompt_version": PROMPT_VERSION},
        "loop1_capture": loop1, "loop2_brief": loop2,
        "betterbriefs_scorecard": scorecard,
    }
    # Loops 3–7 (RAG) only when explicitly enabled — key is omitted otherwise, so
    # output is byte-for-byte identical to a Loops 1–2 run.
    if loops37:
        out["loops3_7"] = loops_3_7(loop2, fields)
    if golden:
        gb = extract_golden_brief(text)
        if gb:
            out["loop2_golden"] = gb

    # Loop 4+5: fill insight + desired_response from IPA precedents + playbooks.
    # Only runs when loops37 ran successfully AND golden extraction produced fields.
    if loops37 and out.get("loops3_7", {}).get("enabled") and out.get("loop2_golden"):
        gf = out["loop2_golden"].setdefault("fields", {})
        golden_schema = json.loads((HERE / "golden-brief" / "golden_brief.schema.json").read_text())
        # Generates insight/smp/rtb/desired_response from the brief + retrieved IPA
        # precedent, schema-driven and rubric-gated. Mutates gf in place; never
        # overwrites a client_stated field. Failures become open questions.
        _fills, gen_open_qs = fill_derivable_fields(gf, out["loops3_7"], golden_schema, brief_text=text)
        if gen_open_qs:
            out["loop2_golden"]["generation_open_questions"] = gen_open_qs
            out["loop2_brief"].setdefault("open_questions", []).extend(gen_open_qs)
    # Snapshot the LLM call ledger so optimisation work is measured per run.
    out["meta"]["llm_stats"] = {**_LLM_STATS,
                                "wall_seconds": round((dt.datetime.now() - _t_run0).total_seconds(), 1)}
    return out


def main():
    ap = argparse.ArgumentParser(description="Briefing tool MVP — Loops 1 & 2")
    ap.add_argument("brief", nargs="?", default=None,
                    help="path to brief (.txt/.md/.docx/.pdf/.eml or an image .png/.jpg), "
                         "or '-' to read pasted text from stdin")
    ap.add_argument("--text", default=None,
                    help="brief text inline (e.g. a copy-pasted email) instead of a file")
    ap.add_argument("--attach", action="append", default=[], metavar="FILE",
                    help="supplementary file(s) folded in as context (e.g. brand guidelines); "
                         "any supported type incl. images/PDF. Repeatable.")
    ap.add_argument("--format", default="md",
                    help="output formats, comma-separated: md,docx,pdf (default: md). "
                         "JSON is always written.")
    ap.add_argument("--out", default=None, help="output dir (default: outputs/<name>)")
    ap.add_argument("--client")
    ap.add_argument("--project")
    ap.add_argument("--provider", help="nim | openai | ollama | anthropic (else auto-detect)")
    ap.add_argument("--model", help="override model id (e.g. nvidia/llama-3.1-nemotron-70b-instruct)")
    ap.add_argument("--check", action="store_true",
                    help="connectivity check: list available models (esp. Nemotron) and exit")
    ap.add_argument("--loops37", action="store_true",
                    help="also run Loops 3–7 (RAG-grounded strategy from rag/index). "
                         "Off by default; needs a built index (cd rag && ./build_rag.sh).")
    ap.add_argument("--golden", action="store_true",
                    help="run schema-grounded Golden Brief extraction pass (loop2_golden).")
    args = ap.parse_args()
    loops37 = args.loops37 or os.environ.get("BRIEF_LOOPS37", "").lower() in ("1", "true", "yes")
    # golden always runs when loops37 is on — insight fill (Loop 4) depends on it
    golden = args.golden or loops37 or os.environ.get("BRIEF_GOLDEN", "").lower() in ("1", "true", "yes")

    if args.provider:
        os.environ["BRIEF_PROVIDER"] = args.provider
    if args.model:
        os.environ["BRIEF_MODEL"] = args.model

    if args.check:
        prov = resolve_provider() or "nim"
        try:
            ids = list_models(prov)
            nem = [m for m in ids if "nemotron" in m.lower()]
            print(f"✓ reachable via '{prov}'. {len(ids)} models. Nemotron ids:")
            print("\n".join(f"  - {m}" for m in nem[:20]) or "  (none found)")
        except Exception as e:
            print(f"✗ can't reach '{prov}': {e.__class__.__name__}: {e}")
        return

    # Resolve the input: inline --text, '-'/piped stdin (pasted email), or a file.
    path = None
    raw_text = None
    source_name = None
    if args.text is not None:
        raw_text = ingest_email_text(args.text)
        source_name = "pasted-email"
    elif args.brief == "-" or (args.brief is None and not sys.stdin.isatty()):
        raw_text = ingest_email_text(sys.stdin.read())
        source_name = "pasted-email"
        if not raw_text.strip():
            sys.exit("No input on stdin. Paste the email then Ctrl-D, or pass a file / --text.")
    elif args.brief:
        path = Path(args.brief).expanduser().resolve()
        if not path.exists():
            sys.exit(f"File not found: {path}")
    else:
        sys.exit("Give a brief: a file path, '-' for stdin, or --text \"...\".")

    # Fold any --attach files (brand guidelines, etc.) into the brief as context.
    if args.attach:
        if raw_text is None:                      # ingest the primary file here so we can append
            raw_text, _ = ingest(path)
            source_name = source_name or path.name
            path = None                           # raw_text now carries it; don't re-ingest in run()
        for att in args.attach:
            ap_path = Path(att).expanduser().resolve()
            if not ap_path.exists():
                sys.exit(f"Attachment not found: {ap_path}")
            print(f"[i] attaching context: {ap_path.name}", file=sys.stderr)
            atext, _ = ingest(ap_path)
            raw_text += (f"\n\n===== ATTACHMENT: {ap_path.name} "
                         f"(supporting context — e.g. brand guidelines) =====\n{atext}")

    brief = run(path, args.client, args.project, loops37=loops37, golden=golden,
                raw_text=raw_text, source_name=source_name)

    stem = source_name or (path.stem if path else "brief")
    name = re.sub(r"[^a-z0-9]+", "-", (args.project or stem).lower()).strip("-")
    out_dir = Path(args.out).resolve() if args.out else (HERE / "outputs" / name)
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "brief_object.json").write_text(json.dumps(brief, indent=2), encoding="utf-8")
    # The deliverable: only the final brief. PDF/DOCX derive from this.
    md_text = render_client_brief(brief)
    md_file = out_dir / "client_brief.md"
    md_file.write_text(md_text, encoding="utf-8")
    # The machinery (capture, ledger, reviews, scorecard, loop narratives) — for the team, not the client.
    (out_dir / "review.md").write_text(render_markdown(brief), encoding="utf-8")
    rich = write_rich_formats(md_text, md_file, args.format.split(","))

    led = brief["loop1_capture"]["no_loss_ledger"]
    print(f"✓ {source_name or path.name}  [mode: {brief['meta']['extraction_mode']}]")
    print(f"  Loop 1 no-loss: {led['coverage_pct']}% "
          f"({led['mapped_segments']}/{led['total_segments']})  "
          f"review: {'pass' if brief['loop1_capture']['review']['passed'] else 'needs pass'}")
    print(f"  Loop 2 open questions: {len(brief['loop2_brief']['open_questions'])}  "
          f"review: {'pass' if brief['loop2_brief']['review']['passed'] else 'gaps'}")
    sc = brief["betterbriefs_scorecard"]
    verdicts = [d["verdict"] for d in sc["dimensions"]]
    sm = sc["single_mindedness"]
    print(f"  BetterBriefs scorecard ({sc['mode']}): "
          f"{verdicts.count('pass')} pass / {verdicts.count('vague')} vague / "
          f"{verdicts.count('missing')} missing"
          + (f"  ⚠️ split into {len(sm['split_into'])} briefs"
             if sm["verdict"] == "multiple" else ""))
    s37 = brief.get("loops3_7")
    if s37:
        if s37.get("enabled"):
            print(f"  Loops 3–7 ({s37['synthesis_mode']}): intent={s37['intent']}, "
                  f"{len(s37['sources_used'])} playbook sections cited")
        else:
            print(f"  Loops 3–7: skipped — {s37.get('reason', '')}")
    if brief.get("loop2_golden"):
        gf = brief["loop2_golden"].get("fields", {})
        filled = sum(1 for v in gf.values() if isinstance(v, dict) and v.get("source") != "missing")
        loop4_ran = "insight" in gf and (gf["insight"] or {}).get("method") == "loop4_fill"
        loop5_ran = "desired_response" in gf and (gf.get("desired_response") or {}).get("method") == "loop5_fill"
        print(f"  Golden Brief: {filled}/{len(gf)} fields filled"
              + (" · insight filled (loop4)" if loop4_ran else "")
              + (" · desired_response filled (loop5)" if loop5_ran else ""))
    formats_written = ["json", "md"] + rich
    print(f"  Output -> {out_dir}  [{', '.join(formats_written)}]")
    ls = brief["meta"].get("llm_stats") or {}
    if ls.get("calls"):
        tok = (f"{ls['prompt_tokens']}+{ls['completion_tokens']} tok"
               if ls.get("prompt_tokens") else f"{ls['input_chars']}+{ls['output_chars']} chars")
        print(f"  LLM: {ls.get('logical_calls', '?')} calls ({ls['calls']} link attempts) · {tok} · "
              f"retries {ls['retries']} · rate-limited {ls['rate_limited']} · {ls.get('wall_seconds', '?')}s")


if __name__ == "__main__":
    main()
